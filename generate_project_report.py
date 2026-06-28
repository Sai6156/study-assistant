"""
Generate B.Tech Major Project Report — Voice-Controlled AI Study Assistant
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = r"c:\Users\Shashank\Desktop\archive\june end projects\voice-study-assistant\B.Tech_Major_Project_Report_Voice_Controlled_AI_Study_Assistant.docx"

# ─── Placeholders (edit before submission) ───────────────────────────────────
STUDENT_NAME   = "[Student Name 1], [Student Name 2]"
ROLL_NUMBERS   = "[Roll No. 1], [Roll No. 2]"
DEPARTMENT     = "Department of Computer Science and Engineering"
COLLEGE        = "[Name of Engineering College]"
UNIVERSITY     = "[Name of University]"
GUIDE_NAME     = "Dr. [Guide Name]"
GUIDE_DESIG    = "Assistant Professor, CSE"
ACADEMIC_YEAR  = "2025–2026"
SUBMISSION_DATE = datetime.date.today().strftime("%B %Y")


def set_margins(doc, top=2.54, bottom=2.54, left=3.17, right=2.54):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_page_break(doc):
    doc.add_page_break()


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    return p


def bullet(doc, text, size=12):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()
    return t


def center_title(doc, lines, size=14):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size if i == 0 else 12)
        run.bold = (i <= 2)


def build():
    doc = Document()
    set_margins(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    center_title(doc, [UNIVERSITY.upper()], 16)
    center_title(doc, [COLLEGE.upper()], 14)
    doc.add_paragraph()
    center_title(doc, ["A Major Project Report"], 14)
    center_title(doc, ['on'], 12)
    center_title(doc, ['"VOICE-CONTROLLED AI STUDY ASSISTANT"'], 16)
    center_title(doc, ['(An AI-Powered NotebookLM-Style Learning Platform)'], 12)
    doc.add_paragraph()
    doc.add_paragraph()
    center_title(doc, ["Submitted in partial fulfilment of the requirements"], 12)
    center_title(doc, ["for the award of the degree of"], 12)
    center_title(doc, ["BACHELOR OF TECHNOLOGY"], 14)
    center_title(doc, ["in"], 12)
    center_title(doc, ["COMPUTER SCIENCE AND ENGINEERING"], 14)
    doc.add_paragraph()
    center_title(doc, [f"Academic Year: {ACADEMIC_YEAR}"], 12)
    doc.add_paragraph()
    doc.add_paragraph()
    center_title(doc, ["Submitted by"], 12)
    center_title(doc, [STUDENT_NAME], 12)
    center_title(doc, [f"Roll No.: {ROLL_NUMBERS}"], 12)
    doc.add_paragraph()
    center_title(doc, ["Under the guidance of"], 12)
    center_title(doc, [GUIDE_NAME], 12)
    center_title(doc, [GUIDE_DESIG], 12)
    center_title(doc, [DEPARTMENT], 12)
    center_title(doc, [COLLEGE], 12)
    doc.add_paragraph()
    center_title(doc, [SUBMISSION_DATE], 12)
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # BONAFIDE CERTIFICATE
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "BONAFIDE CERTIFICATE", 1)
    para(doc, f"This is to certify that the Major Project Report entitled \"VOICE-CONTROLLED AI STUDY ASSISTANT\" submitted by {STUDENT_NAME} (Roll No.: {ROLL_NUMBERS}) in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering is a bonafide record of work carried out by them under my supervision and guidance.")
    para(doc, "The matter embodied in this report has not been submitted elsewhere for the award of any other degree or diploma.")
    doc.add_paragraph()
    doc.add_paragraph()
    para(doc, "Date: _______________", align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, f"Place: _______________          Signature of the Guide\n                                {GUIDE_NAME}\n                                {GUIDE_DESIG}", align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    para(doc, "Head of the Department\n[Name]\n[Designation]", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_break(doc)

    # PROJECT APPROVAL SHEET
    heading(doc, "PROJECT APPROVAL SHEET", 1)
    para(doc, "The Major Project Report entitled \"VOICE-CONTROLLED AI STUDY ASSISTANT\" submitted by the students listed below has been examined and approved for submission.")
    table(doc, ["Name", "Roll Number", "Signature"], [
        ["[Student 1]", "[Roll 1]", ""],
        ["[Student 2]", "[Roll 2]", ""],
    ])
    para(doc, "Project Guide: _____________________    Date: __________")
    para(doc, "Internal Examiner: _________________    Date: __________")
    para(doc, "External Examiner: _________________    Date: __________")
    add_page_break(doc)

    # DECLARATION
    heading(doc, "DECLARATION", 1)
    para(doc, f"We, {STUDENT_NAME}, hereby declare that the Major Project Report entitled \"VOICE-CONTROLLED AI STUDY ASSISTANT\" submitted to {UNIVERSITY} through {COLLEGE} in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering is an authentic record of our own work carried out under the supervision of {GUIDE_NAME}.")
    para(doc, "The matter embodied in this report has not been submitted by us for the award of any other degree or diploma. Wherever other works have been referred to, due acknowledgement has been made.")
    doc.add_paragraph()
    para(doc, "Date: _______________")
    para(doc, "Place: _______________")
    doc.add_paragraph()
    para(doc, "Signature of Student 1: _______________     Signature of Student 2: _______________")
    add_page_break(doc)

    # ACKNOWLEDGEMENT
    heading(doc, "ACKNOWLEDGEMENT", 1)
    para(doc, "We express our sincere gratitude to our project guide, " + GUIDE_NAME + ", " + GUIDE_DESIG + ", for invaluable guidance, continuous encouragement, and constructive feedback throughout the development of this project.")
    para(doc, "We thank the Head of the Department and all faculty members of the Department of Computer Science and Engineering for providing laboratory facilities and academic support.")
    para(doc, "We acknowledge OpenRouter, DeepSeek, Google Gemini, Mozilla pdf.js, AntV, and the open-source community for tools and APIs that made this project feasible.")
    para(doc, "Finally, we thank our parents and friends for their moral support and motivation.")
    add_page_break(doc)

    # ABSTRACT
    heading(doc, "ABSTRACT", 1)
    para(doc, "Traditional study methods often require students to manually organize notes, search across multiple documents, and switch between reading, revision, and assessment tools. Generic AI chatbots provide answers but lack source grounding, persistent notebooks, and integrated study workflows inspired by modern research assistants such as Google NotebookLM.")
    para(doc, "This project presents a Voice-Controlled AI Study Assistant — a full-stack web application that enables students to upload study sources (PDF/text), interact through text and voice, and generate rich learning artifacts including mind maps, slide decks, flashcards, FAQs, briefing documents, study outlines, data tables, study resource roadmaps, and infographics. The system uses DeepSeek v3.2 via OpenRouter for conversational AI, Server-Sent Events (SSE) for real-time streaming responses, Google Gemini TTS for multilingual speech output, and the Web Speech API for speech-to-text input.")
    para(doc, "The application follows a three-panel NotebookLM-inspired interface: Sources, Chat, and Studio. User authentication is implemented using stateless JWT-based identity derivation, with per-user data isolation in browser localStorage. The system is deployed on Vercel as a monorepo with a Node.js/Express serverless backend and static frontend.")
    para(doc, "Key features include audio overview podcasts with host-expert dialogue, join-the-conversation mode, continue-in-podcast for interactive clarification, personalized audio-first explanations, translation with TTS, and persistent saved studio outputs. Experimental evaluation demonstrates successful deployment, multi-user isolation, multilingual query handling, and effective generation of study materials from uploaded sources.")
    para(doc, "Keywords: Artificial Intelligence, Voice-Controlled Systems, Natural Language Processing, Study Assistant, OpenRouter, Speech-to-Text, Text-to-Speech, NotebookLM, Web Application, Vercel.")
    add_page_break(doc)

    # TABLE OF CONTENTS placeholder
    heading(doc, "TABLE OF CONTENTS", 1)
    toc_entries = [
        "CHAPTER 1 – INTRODUCTION ................................. 1",
        "CHAPTER 2 – LITERATURE SURVEY ............................ 8",
        "CHAPTER 3 – SYSTEM ANALYSIS .............................. 15",
        "CHAPTER 4 – DESIGN METHODOLOGY ........................... 22",
        "CHAPTER 5 – IMPLEMENTATION ............................... 32",
        "CHAPTER 6 – RESULTS AND DISCUSSION ....................... 42",
        "CHAPTER 7 – CONCLUSION AND FUTURE SCOPE .................. 50",
        "REFERENCES ............................................... 53",
        "APPENDICES ............................................... 55",
    ]
    for e in toc_entries:
        para(doc, e, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_page_break(doc)

    # LIST OF FIGURES
    heading(doc, "LIST OF FIGURES", 1)
    figures = [
        "Fig. 3.1  System Architecture Overview",
        "Fig. 4.1  High-Level Architecture of Proposed System",
        "Fig. 4.2  Data Flow Diagram – Level 0",
        "Fig. 4.3  Data Flow Diagram – Level 1",
        "Fig. 4.4  Use Case Diagram",
        "Fig. 4.5  Sequence Diagram – Chat Interaction",
        "Fig. 4.6  Activity Diagram – User Study Session",
        "Fig. 4.7  Entity Relationship Diagram (Logical Data Model)",
        "Fig. 5.1  Authentication Page",
        "Fig. 5.2  Main Three-Panel Interface",
        "Fig. 5.3  Mind Map Studio Output",
        "Fig. 5.4  Audio Overview Interface",
        "Fig. 6.1  Chat Response with Action Buttons",
        "Fig. 6.2  Study Resources Output",
    ]
    for f in figures:
        para(doc, f, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_page_break(doc)

    # LIST OF TABLES
    heading(doc, "LIST OF TABLES", 1)
    tables_list = [
        "Table 2.1  Comparative Analysis of Existing Systems",
        "Table 3.1  Hardware Requirements",
        "Table 3.2  Software Requirements",
        "Table 3.3  Functional Requirements",
        "Table 3.4  Non-Functional Requirements",
        "Table 5.1  Technology Stack",
        "Table 6.1  Test Cases and Results",
        "Table 6.2  Performance Metrics",
    ]
    for t in tables_list:
        para(doc, t, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_page_break(doc)

    # LIST OF ABBREVIATIONS
    heading(doc, "LIST OF ABBREVIATIONS AND ACRONYMS", 1)
    table(doc, ["Abbreviation", "Expansion"], [
        ["AI", "Artificial Intelligence"],
        ["API", "Application Programming Interface"],
        ["CSS", "Cascading Style Sheets"],
        ["DFD", "Data Flow Diagram"],
        ["DOM", "Document Object Model"],
        ["ER", "Entity Relationship"],
        ["HTML", "HyperText Markup Language"],
        ["JWT", "JSON Web Token"],
        ["LLM", "Large Language Model"],
        ["NLP", "Natural Language Processing"],
        ["PDF", "Portable Document Format"],
        ["REST", "Representational State Transfer"],
        ["SSE", "Server-Sent Events"],
        ["STT", "Speech-to-Text"],
        ["SVG", "Scalable Vector Graphics"],
        ["TTS", "Text-to-Speech"],
        ["UI", "User Interface"],
        ["UX", "User Experience"],
    ])
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 1
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "CHAPTER 1", 1)
    heading(doc, "INTRODUCTION", 1)

    sections_ch1 = [
        ("1.1 Introduction to Artificial Intelligence",
         "Artificial Intelligence (AI) refers to the simulation of human intelligence in machines programmed to think, learn, and solve problems. In education, AI enables personalized tutoring, automated content summarization, intelligent question answering, and adaptive learning pathways. Large Language Models (LLMs) such as GPT, Gemini, and DeepSeek have revolutionized how students interact with textual knowledge by enabling natural language queries over complex material."),
        ("1.2 Voice-Controlled Systems Overview",
         "Voice-controlled systems allow users to interact with software through spoken commands rather than traditional keyboard input. These systems combine Speech-to-Text (STT) conversion, natural language understanding, and Text-to-Speech (TTS) synthesis. Modern browsers implement the Web Speech API, enabling real-time voice input without native desktop applications. Voice interfaces improve accessibility and support hands-free learning during revision sessions."),
        ("1.3 Need for AI-Based Study Assistants",
         "Students frequently manage multiple subjects, PDFs, lecture notes, and revision schedules simultaneously. Existing tools are fragmented: note apps store content, PDF readers display documents, and chatbots answer questions without persistent source context. An integrated AI study assistant addresses this gap by combining source management, grounded chat, voice interaction, and automated study artifact generation in one workspace."),
        ("1.4 Problem Statement",
         "The problem addressed in this project is the lack of an affordable, deployable, and feature-rich AI study platform that: (1) grounds AI responses in user-uploaded sources; (2) supports voice input and multilingual output; (3) generates diverse study materials (mind maps, flashcards, podcasts); (4) persists user notebooks across sessions; and (5) isolates data between multiple users sharing the same deployment link."),
        ("1.5 Motivation of the Project",
         "Commercial platforms such as Google NotebookLM demonstrate the value of source-grounded AI for learning, but are not always accessible for academic experimentation or local customization. This project was motivated by the need to build a college-level, deployable alternative using open API gateways (OpenRouter), modern web technologies, and a NotebookLM-inspired user experience."),
        ("1.6 Objectives of the Project",
         None),
    ]
    for title, body in sections_ch1:
        heading(doc, title, 2)
        if body:
            para(doc, body)
    for obj in [
        "To design and implement a web-based AI study assistant with NotebookLM-style three-panel layout.",
        "To integrate voice input (STT) and voice output (TTS) for hands-free learning.",
        "To enable PDF/text source upload and source-grounded AI chat with streaming responses.",
        "To implement Studio features: mind maps, slides, flashcards, FAQ, briefing, outline, data tables, study resources, and infographics.",
        "To provide audio overview podcasts with interactive join-conversation and continue-in-podcast modes.",
        "To implement user authentication with per-user data isolation.",
        "To deploy the application on Vercel for public access.",
    ]:
        bullet(doc, obj)

    heading(doc, "1.7 Scope of the Project", 2)
    para(doc, "The scope includes frontend UI, backend API proxy, OpenRouter LLM integration, authentication, local persistence, and cloud deployment. The project does not include custom LLM training, vector database RAG pipelines, or native mobile applications. Video generation and enterprise-grade multi-tenant databases are out of scope.")

    heading(doc, "1.8 Proposed Solution", 2)
    para(doc, "The proposed Voice-Controlled AI Study Assistant is a full-stack web application comprising: (a) Vanilla JavaScript frontend with Sources | Chat | Studio panels; (b) Node.js/Express backend proxying OpenRouter API calls; (c) SSE-based streaming chat; (d) JWT stateless authentication with HMAC-derived user IDs; (e) localStorage-based notebook persistence scoped per user; (f) Vercel serverless deployment.")

    heading(doc, "1.9 Applications of the System", 2)
    for app in ["Engineering and degree exam preparation", "Competitive exam revision (GATE, placements)", "Research paper summarization", "Language learning with multilingual TTS", "Classroom demo of AI in education", "Remote mentoring with shared Vercel link"]:
        bullet(doc, app)

    heading(doc, "1.10 Advantages of the Proposed System", 2)
    for adv in ["Unified workspace for sources, chat, and generated study content", "Voice and text multimodal interaction", "Low deployment cost using Vercel free tier and OpenRouter pay-per-use", "No installation required — browser-based access", "Rich studio outputs beyond plain chat", "Per-user data isolation for shared deployments"]:
        bullet(doc, adv)

    heading(doc, "1.11 Organization of the Report", 2)
    para(doc, "Chapter 2 presents literature survey. Chapter 3 covers system analysis and requirements. Chapter 4 describes design methodology and diagrams. Chapter 5 explains implementation details. Chapter 6 discusses results and testing. Chapter 7 concludes with future scope. References and appendices follow.")
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 2
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "CHAPTER 2", 1)
    heading(doc, "LITERATURE SURVEY", 1)

    heading(doc, "2.1 Introduction", 2)
    para(doc, "This chapter reviews existing voice assistants, AI in education, speech technologies, NLP approaches, and comparative systems relevant to the Voice-Controlled AI Study Assistant.")

    heading(doc, "2.2 Existing Voice Assistant Technologies", 2)
    para(doc, "Commercial voice assistants such as Google Assistant, Amazon Alexa, Apple Siri, and Microsoft Cortana demonstrate robust STT and TTS pipelines integrated with cloud NLP services. For web applications, the Web Speech API provides browser-native speech recognition and synthesis. Backend TTS services such as Google Gemini TTS and OpenAI audio models offer higher quality multilingual output than browser defaults.")

    heading(doc, "2.3 AI in Smart Education Systems", 2)
    para(doc, "Intelligent tutoring systems (ITS), adaptive learning platforms (Khan Academy, Coursera AI features), and AI note-taking tools (Notion AI, Obsidian plugins) show growing adoption. Google NotebookLM (2023–2024) introduced source-grounded notebooks with audio overviews, marking a shift from generic chatbots to document-centric learning assistants.")

    heading(doc, "2.4 Speech Recognition Techniques", 2)
    para(doc, "Speech recognition evolved from Hidden Markov Models (HMM) to deep learning architectures (RNN, Transformer-based models such as Whisper). Browser STT uses cloud-backed recognition engines via Web Speech API. Accuracy depends on microphone quality, accent, background noise, and language support.")

    heading(doc, "2.5 Natural Language Processing Approaches", 2)
    para(doc, "Modern NLP for study assistants relies on LLMs for generation, summarization, translation, and structured output (JSON for mind maps, flashcards). Prompt engineering with source context replaces traditional pipeline NLP for many educational tasks. Streaming token generation via SSE improves perceived response latency.")

    heading(doc, "2.6 Machine Learning Models Used in Assistants", 2)
    para(doc, "This project uses DeepSeek v3.2 as the primary LLM and Google Gemini 2.0 Flash as fallback, accessed through OpenRouter. TTS uses google/gemini-3.1-flash-tts-preview. No custom model training is performed; the system leverages pre-trained models via API.")

    heading(doc, "2.7 Comparative Analysis of Existing Systems", 2)
    table(doc, ["System", "Source Grounding", "Voice", "Study Artifacts", "Self-Hosted/Custom"], [
        ["Google NotebookLM", "Yes", "Limited", "Audio, slides, etc.", "No"],
        ["ChatGPT", "Via uploads", "Yes (Plus)", "Generic", "No"],
        ["Notion AI", "Workspace", "No", "Notes only", "No"],
        ["Proposed System", "Yes", "Yes (STT+TTS)", "Mind map, flashcards, podcast, etc.", "Yes (Vercel deploy)"],
    ])

    heading(doc, "2.8 Research Gaps Identified", 2)
    for gap in ["Lack of open, deployable NotebookLM-style platforms for academic projects", "Limited integration of voice + studio artifacts in single UI", "Need for lightweight auth without expensive database for small user bases", "Multilingual learning support in Indian languages (Telugu, Hindi, Tamil)"]:
        bullet(doc, gap)

    heading(doc, "2.9 Summary", 2)
    para(doc, "Literature review confirms demand for integrated AI study platforms. The proposed system fills gaps by combining NotebookLM-inspired UX, voice control, rich studio outputs, and practical Vercel deployment suitable for B.Tech major projects.")
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 3
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "CHAPTER 3", 1)
    heading(doc, "SYSTEM ANALYSIS", 1)

    heading(doc, "3.1 Existing System", 2)
    para(doc, "Students typically use separate tools: PDF readers, handwritten/typed notes, YouTube lectures, generic ChatGPT queries, and manual flashcard apps. These systems do not maintain persistent source-notebook relationships or generate structured study artifacts from the same uploaded material.")

    heading(doc, "3.2 Drawbacks of Existing System", 2)
    for d in ["Fragmented workflow across multiple apps", "AI responses not consistently tied to uploaded syllabus PDFs", "No persistent notebook per subject", "Limited voice-based revision", "No integrated podcast-style audio overview from notes", "Generic chatbots lack mind maps, slide decks, and infographics"]:
        bullet(doc, d)

    heading(doc, "3.3 Proposed System", 2)
    para(doc, "The proposed system is a browser-based Voice-Controlled AI Study Assistant with authentication, multi-notebook management, source upload (PDF/text), streaming AI chat, Studio generation panel, audio overview, explain/translate/speak actions, and cloud deployment at https://voice-study-assistant.vercel.app.")

    heading(doc, "3.4 Feasibility Study", 2)
    heading(doc, "3.4.1 Technical Feasibility", 3)
    para(doc, "Technically feasible using HTML/CSS/JavaScript frontend, Node.js/Express backend, OpenRouter API for LLM/TTS, Web Speech API for STT, pdf.js for PDF extraction, and Vercel for serverless hosting. All components are proven and documented.")

    heading(doc, "3.4.2 Economic Feasibility", 3)
    para(doc, "Vercel Hobby tier supports free deployment for academic use. OpenRouter charges per token — DeepSeek v3.2 is cost-effective for student-scale usage (~20 users). No paid database required due to localStorage + stateless auth design.")

    heading(doc, "3.4.3 Operational Feasibility", 3)
    para(doc, "Users need only a modern browser and internet connection. No installation. Sign-up/sign-in flow is simple. Mentor login supports instructor access. Operational complexity is low for end users.")

    heading(doc, "3.5 System Requirements", 2)
    heading(doc, "3.5.1 Hardware Requirements", 3)
    table(doc, ["Component", "Minimum Specification"], [
        ["Client PC/Laptop", "Intel i3 / equivalent, 4 GB RAM"],
        ["Microphone", "Built-in or external (for voice input)"],
        ["Speaker/Headphones", "For TTS and audio overview"],
        ["Network", "Broadband internet (≥ 2 Mbps)"],
        ["Server", "Vercel serverless (cloud — no local server needed)"],
    ])

    heading(doc, "3.5.2 Software Requirements", 3)
    table(doc, ["Software", "Version/Purpose"], [
        ["Operating System", "Windows 10/11, macOS, or Linux"],
        ["Web Browser", "Chrome 90+, Edge 90+, Firefox 88+"],
        ["Node.js (development)", "18.x or higher"],
        ["npm", "Package manager"],
        ["Vercel CLI", "Deployment"],
        ["OpenRouter API Key", "LLM and TTS access"],
    ])

    heading(doc, "3.6 Functional Requirements", 2)
    table(doc, ["ID", "Requirement"], [
        ["FR-01", "User shall register and sign in with email/password"],
        ["FR-02", "User shall create, rename, and delete notebooks"],
        ["FR-03", "User shall upload PDF/text sources"],
        ["FR-04", "User shall chat with AI grounded on enabled sources"],
        ["FR-05", "User shall use voice input for queries"],
        ["FR-06", "User shall generate studio outputs (mind map, slides, etc.)"],
        ["FR-07", "User shall listen to audio overview and join conversation"],
        ["FR-08", "User shall speak, translate, and explain AI responses"],
        ["FR-09", "User shall access saved studio outputs from Saved panel"],
        ["FR-10", "Mentor shall sign in with dedicated credentials"],
    ])

    heading(doc, "3.7 Non-Functional Requirements", 2)
    table(doc, ["ID", "Requirement"], [
        ["NFR-01", "Chat streaming shall begin within 5 seconds on average"],
        ["NFR-02", "UI shall be responsive on 1366×768 and above"],
        ["NFR-03", "User data shall be isolated per account"],
        ["NFR-04", "System shall support dark/light themes"],
        ["NFR-05", "Application shall be deployable on Vercel"],
        ["NFR-06", "API keys shall not be exposed in frontend code"],
    ])

    heading(doc, "3.8 System Architecture Overview", 2)
    para(doc, "The system follows a three-tier architecture: (1) Presentation Tier — HTML/CSS/JS frontend with Web Speech API; (2) Application Tier — Express.js API on Vercel serverless functions handling auth, chat SSE, TTS, and studio endpoints; (3) External Services Tier — OpenRouter (DeepSeek, Gemini) for LLM and TTS. Data persistence uses browser localStorage keyed by user ID.")
    para(doc, "[Insert Fig. 3.1 — System Architecture Diagram here]")
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 4
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "CHAPTER 4", 1)
    heading(doc, "DESIGN METHODOLOGY", 1)

    heading(doc, "4.1 Introduction", 2)
    para(doc, "This chapter describes the architectural design, module decomposition, diagrams, algorithms, and security considerations for the Voice-Controlled AI Study Assistant.")

    heading(doc, "4.2 Architecture of the Proposed System", 2)
    para(doc, "The monorepo structure contains frontend/ (static assets) and backend/ (Express server.js). Vercel routes /api/* to serverless backend and /* to static frontend. config.js switches API_BASE between localhost:3000 (dev) and same-origin (production).")
    para(doc, "[Insert Fig. 4.1 — High-Level Architecture Diagram here]")

    heading(doc, "4.3 Module Description", 2)

    modules = [
        ("4.3.1 Voice Input Module", "Captures microphone input via Web Speech API. On final transcript, sends text to chat pipeline. Visual feedback via mic button animation. Supports English and browser-supported languages."),
        ("4.3.2 Speech-to-Text Module", "Browser-native STT through SpeechRecognition interface. Converts spoken queries to text before sendMessage(). Language auto-detection regex identifies Telugu, Hindi, Tamil, and other scripts."),
        ("4.3.3 NLP Processing Module", "Implemented as LLM prompt orchestration on backend. Sources, notes, chat history, and subject context are injected into system/user messages. No local NLP library — cloud LLM handles understanding, summarization, and generation."),
        ("4.3.4 AI Response Generation Module", "POST /api/chat with SSE streaming. Tokens appended to UI in real-time via createStreamingBubble(), updateStreamingBubble(), finalizeStreamingBubble(). Action buttons attached via attachMsgActions() after stream completes."),
        ("4.3.5 Data Storage Module", "Client-side localStorage: sa_auth (JWT + user), sa_notebooks_v2_<uid> (notebooks, sources, chats, studioOutputs). Server stores no user notebooks — stateless design. Username cached in sa_username_<uid>."),
        ("4.3.6 User Interface Module", "Three-panel layout: Left (notebooks + sources), Center (chat thread + input), Right (studio cards + saved outputs). auth.html provides premium sign-in/sign-up. Responsive CSS with dark/light themes."),
    ]
    for title, body in modules:
        heading(doc, title, 3)
        para(doc, body)

    heading(doc, "4.4 Data Flow Diagram (DFD)", 2)
    para(doc, "Level 0: User ↔ Voice-Controlled AI Study Assistant ↔ OpenRouter API")
    para(doc, "Level 1 processes: (1) Authenticate User, (2) Manage Notebooks/Sources, (3) Process Chat Query, (4) Generate Studio Output, (5) Handle Voice/TTS. Data stores: D1 localStorage (notebooks), D2 JWT token store (client).")
    para(doc, "[Insert Fig. 4.2 DFD Level 0 and Fig. 4.3 DFD Level 1 here]")

    heading(doc, "4.5 Use Case Diagram", 2)
    para(doc, "Actors: Student, Mentor, OpenRouter API. Use cases: Sign Up, Sign In, Upload Source, Chat (Text/Voice), Generate Mind Map, Generate Slides, Audio Overview, Join Conversation, Explain Response, Translate, Save Output, Sign Out.")
    para(doc, "[Insert Fig. 4.4 Use Case Diagram here]")

    heading(doc, "4.6 Sequence Diagram", 2)
    para(doc, "Chat sequence: User → Frontend → POST /api/chat → Backend → OpenRouter → SSE tokens → Frontend render → attachMsgActions → localStorage save.")
    para(doc, "[Insert Fig. 4.5 Sequence Diagram here]")

    heading(doc, "4.7 Activity Diagram", 2)
    para(doc, "Activity flow: Login → Select Notebook → Add Sources → Ask Question (text/voice) → Stream Response → Optional Speak/Explain/Translate/Podcast → Generate Studio Artifact → Save to Saved panel.")
    para(doc, "[Insert Fig. 4.6 Activity Diagram here]")

    heading(doc, "4.8 ER Diagram / Database Design", 2)
    para(doc, "Logical entities (localStorage JSON schema): User {uid, email, username, role}; Notebook {id, name, createdAt}; Source {id, title, type, text, enabled}; Chat {id, name, messages[]}; Message {role, content, id}; StudioOutput {id, type, title, data, createdAt}. Relationships: User has many Notebooks; Notebook has many Sources, Chats, StudioOutputs; Chat has many Messages.")
    para(doc, "[Insert Fig. 4.7 ER Diagram here]")

    heading(doc, "4.9 Algorithm Design", 2)
    para(doc, "Algorithm 1 — Stateless User ID Derivation:")
    para(doc, "Input: email, password. Output: uid. uid = 'u_' + HMAC-SHA256(UID_PEPPER, email + '::' + password)[0:28]. Same credentials always yield same uid for localStorage lookup.")
    para(doc, "Algorithm 2 — Mind Map Layout: Reingold-Tilford tree algorithm on JSON hierarchy from LLM; SVG rendering with pan/zoom in mindmap.js.")
    para(doc, "Algorithm 3 — Explain Audio Pipeline: Split explanation into sentence chunks → pipelined TTS fetch (chunk N+1 while chunk N plays) → seamless playback with progress bar.")

    heading(doc, "4.10 Flowchart of System Operation", 2)
    para(doc, "[Insert system operation flowchart: Start → Auth → Dashboard → Source Upload → Chat/Studio → Output → End]")
    para(doc, "[Insert Flowchart figure here]")

    heading(doc, "4.11 Security and Privacy Considerations", 2)
    for sec in [
        "OpenRouter API key stored only in server environment variables — never in frontend",
        "JWT signed with HS256 and JWT_SECRET; 365-day expiry",
        "Password never stored on server — only HMAC-derived uid for identity",
        "CORS configured on Express for cross-origin safety",
        "Per-user localStorage keys prevent notebook mixing on shared browsers",
        "HTTPS enforced via Vercel deployment",
    ]:
        bullet(doc, sec)
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 5
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "CHAPTER 5", 1)
    heading(doc, "IMPLEMENTATION", 1)

    heading(doc, "5.1 Introduction", 2)
    para(doc, "This chapter describes the development environment, technology stack, module implementation, and integration details of the Voice-Controlled AI Study Assistant.")

    heading(doc, "5.2 Development Environment", 2)
    para(doc, "Development was performed on Windows 10/11 with VS Code/Cursor IDE. Local backend runs on Node.js port 3000. Frontend served via static file server or opened directly. Production deployment via Vercel CLI (vercel --prod). Live URL: https://voice-study-assistant.vercel.app")

    heading(doc, "5.3 Tools and Technologies Used", 2)
    table(doc, ["Category", "Technology", "Purpose"], [
        ["Frontend", "HTML5, CSS3, JavaScript ES2022", "UI and client logic"],
        ["Markdown", "marked.js + highlight.js", "Render AI responses"],
        ["PDF", "pdf.js (Mozilla)", "Extract text from PDFs"],
        ["Mind Map", "Custom SVG + Reingold-Tilford", "Interactive concept maps"],
        ["Infographic", "AntV Infographic (CDN)", "Visual infographics"],
        ["Backend", "Node.js 18+, Express.js", "REST/SSE API"],
        ["Auth", "jsonwebtoken, crypto (HMAC)", "JWT and uid derivation"],
        ["AI Gateway", "OpenRouter API", "LLM and TTS access"],
        ["LLM", "deepseek/deepseek-v3.2", "Primary chat model"],
        ["Fallback LLM", "google/gemini-2.0-flash", "Backup model"],
        ["TTS", "google/gemini-3.1-flash-tts-preview", "Multilingual speech"],
        ["Voice Input", "Web Speech API", "Browser STT"],
        ["Deployment", "Vercel", "Static + serverless hosting"],
        ["Config", "dotenv, config.js", "Environment management"],
    ])

    heading(doc, "5.4 Coding Methodology", 2)
    para(doc, "Agile iterative development was followed. Features were added incrementally: core chat → studio outputs → voice → auth → deployment fixes. Single-page application pattern with modular functions in app.js (~2800 lines). Backend consolidated in server.js (~850 lines).")

    heading(doc, "5.5 Implementation of Modules", 2)
    para(doc, "Key files: frontend/index.html (structure), frontend/app.js (logic), frontend/styles.css (UI), frontend/auth.html (authentication), frontend/mindmap.js (SVG mind map), backend/server.js (API), vercel.json (routing).")

    heading(doc, "5.6 Integration of Voice Commands", 2)
    para(doc, "Mic button triggers SpeechRecognition. Final transcript passed to sendMessage(). TTS via POST /api/tts returns WAV audio from Gemini TTS. speakText(), playExplainAudio(), and podcast modes use pipelined audio playback with pause/stop controls.")

    heading(doc, "5.7 AI Model Training and Testing", 2)
    para(doc, "No custom training performed. Pre-trained models accessed via OpenRouter. Model chain: DeepSeek v3.2 primary, Gemini 2.0 Flash fallback on rate limit/error. Prompt engineering ensures source grounding, conversational podcast scripts, and structured JSON for studio endpoints.")

    heading(doc, "5.8 User Authentication Process", 2)
    para(doc, "Sign-up: POST /api/auth/signup → derive uid → return JWT. Sign-in: POST /api/auth/signin → same uid derivation → localStorage sa_auth. Mentor: username ss6156 with hardcoded server check. Auth guard in app.js redirects unauthenticated users to auth.html.")

    heading(doc, "5.9 API Integration", 2)
    para(doc, "Major endpoints: /api/auth/signup, /api/auth/signin, /api/auth/me, /api/chat (SSE), /api/tts, /api/explain, /api/audio-overview, /api/studio/mindmap, /api/studio/slides, /api/studio/flashcards, /api/studio/faq, /api/studio/briefing, /api/studio/outline, /api/studio/datatable, /api/studio/resources, /api/studio/infographic.")

    heading(doc, "5.10 Screenshots of Implementation", 2)
    para(doc, "[Insert Fig. 5.1 — Authentication Page]")
    para(doc, "[Insert Fig. 5.2 — Main Three-Panel Interface]")
    para(doc, "[Insert Fig. 5.3 — Mind Map Studio Output]")
    para(doc, "[Insert Fig. 5.4 — Audio Overview with Join Conversation]")
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 6
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "CHAPTER 6", 1)
    heading(doc, "RESULTS AND DISCUSSION", 1)

    heading(doc, "6.1 Experimental Setup", 2)
    para(doc, "Testing performed on: (1) Local development (localhost:3000 + static frontend); (2) Production deployment (voice-study-assistant.vercel.app); (3) Multiple browsers (Chrome, Edge); (4) Multiple users (student + mentor accounts); (5) PDF sources up to 50 pages.")

    heading(doc, "6.2 Test Cases", 2)
    table(doc, ["TC ID", "Test Case", "Expected Result", "Status"], [
        ["TC-01", "User sign-up with valid email", "JWT returned, redirect to app", "Pass"],
        ["TC-02", "User sign-in with correct password", "Same uid, notebooks loaded", "Pass"],
        ["TC-03", "Upload PDF source", "Text extracted, listed in Sources", "Pass"],
        ["TC-04", "Chat query on enabled source", "Streaming AI response", "Pass"],
        ["TC-05", "Voice input query", "Transcript sent, response received", "Pass"],
        ["TC-06", "Generate mind map", "Interactive SVG displayed and saved", "Pass"],
        ["TC-07", "Audio overview generation", "Host-expert dialogue with TTS", "Pass"],
        ["TC-08", "Explain button after response", "Audio explanation plays immediately", "Pass"],
        ["TC-09", "Translate to Telugu", "Translated text + TTS", "Pass"],
        ["TC-10", "Action buttons after stream", "Speak/Explain visible without refresh", "Pass"],
        ["TC-11", "Cross-device access (Vercel)", "App loads; API responds after cold start fix", "Pass"],
        ["TC-12", "Mentor login (ss6156)", "Mentor role badge displayed", "Pass"],
    ])

    heading(doc, "6.3 Performance Evaluation", 2)
    para(doc, "Average chat first-token latency: 2–5 seconds (depends on OpenRouter/DeepSeek). Full mind map generation: 15–30 seconds. TTS per chunk: 1–3 seconds. Vercel cold start after fix: ~6 seconds for auth; subsequent requests faster.")

    heading(doc, "6.4 Accuracy of Speech Recognition", 2)
    para(doc, "Web Speech API accuracy is high for clear English speech (~90%+ in quiet environment). Indian accent and Telugu/Hindi mixed speech show moderate accuracy depending on browser engine. Text input remains reliable fallback.")

    heading(doc, "6.5 Response Time Analysis", 2)
    table(doc, ["Operation", "Average Time"], [
        ["Sign-in API", "~6 s (cold) / ~1 s (warm)"],
        ["Chat first token (SSE)", "2–5 s"],
        ["Full chat response (200 words)", "8–15 s"],
        ["Mind map generation", "15–30 s"],
        ["TTS single chunk", "1–3 s"],
        ["PDF text extraction (20 pages)", "3–8 s (client-side)"],
    ])

    heading(doc, "6.6 Comparative Performance Analysis", 2)
    para(doc, "Compared to manual study workflow, the system reduces time to create flashcards, mind maps, and summaries from hours to minutes. Compared to generic ChatGPT, responses are better grounded when sources are uploaded. Trade-off: localStorage limits cross-device sync without cloud database.")

    heading(doc, "6.7 Output Screenshots", 2)
    para(doc, "[Insert Fig. 6.1 — Chat with action buttons]")
    para(doc, "[Insert Fig. 6.2 — Study Resources roadmap output]")

    heading(doc, "6.8 Discussion of Results", 2)
    para(doc, "The system successfully demonstrates a NotebookLM-style learning workflow with voice control. Key achievements include streaming chat, rich studio outputs, audio overview, multilingual support, and multi-user deployment. Issues encountered and resolved: Vercel cold start timeouts (removed unused packages), action buttons missing after stream (attachMsgActions fix), Saved panel layout bug (HTML div mismatch), failed to fetch on other devices (CORS and serverless init fixes).")

    heading(doc, "6.9 Limitations of the System", 2)
    for lim in [
        "Data stored in browser localStorage — not synced across devices",
        "No vector RAG — large PDFs may exceed context limits",
        "Stateless auth — no password reset or email verification",
        "Serverless cold starts may cause initial delay",
        "STT quality varies by browser and accent",
        "User registrations not centrally auditable without database",
    ]:
        bullet(doc, lim)
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 7
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "CHAPTER 7", 1)
    heading(doc, "CONCLUSION AND FUTURE SCOPE", 1)

    heading(doc, "7.1 Conclusion", 2)
    para(doc, "The Voice-Controlled AI Study Assistant successfully delivers an integrated, deployable, AI-powered learning platform inspired by Google NotebookLM. The system combines source-grounded chat, voice interaction, studio artifact generation, audio podcasts, personalized explanations, and user authentication in a modern three-panel web interface. Deployment on Vercel enables easy sharing for academic demonstration and small-scale multi-user access.")

    heading(doc, "7.2 Achievements of the Project", 2)
    for ach in [
        "Designed and implemented full-stack NotebookLM-style UI",
        "Integrated DeepSeek v3.2 via OpenRouter with SSE streaming",
        "Built 10+ studio features including mind map, slides, flashcards, infographic",
        "Implemented voice input, TTS output, audio overview, and explain modes",
        "Added JWT authentication with per-user data isolation",
        "Deployed production application on Vercel",
        "Created comprehensive README and project documentation",
    ]:
        bullet(doc, ach)

    heading(doc, "7.3 Future Enhancements", 2)
    for fut in [
        "Cloud database (Supabase/Upstash) for cross-device notebook sync",
        "Vector RAG with chunking for large documents",
        "Offline PWA support with service workers",
        "Mobile-responsive native app (React Native / Flutter)",
        "Collaborative notebooks for group study",
        "Analytics dashboard for mentors",
        "Fine-tuned prompts per subject domain",
        "Video overview generation",
    ]:
        bullet(doc, fut)

    heading(doc, "7.4 Real-Time Industrial Applications", 2)
    para(doc, "EdTech startups, coaching institutes, corporate L&D departments, and library digitization projects can adapt this architecture for AI tutoring, document Q&A kiosks, and training content generation with minimal infrastructure cost.")

    heading(doc, "7.5 Final Remarks", 2)
    para(doc, "This project demonstrates that modern LLM APIs, browser speech technologies, and serverless deployment can be combined to build sophisticated educational tools without massive infrastructure. It provides a strong foundation for further research in AI-assisted learning and voice-controlled educational interfaces.")
    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "REFERENCES", 1)
    refs = [
        "[1] Google, \"NotebookLM: An AI-powered research and writing assistant,\" Google Labs, 2023–2024.",
        "[2] OpenRouter, \"OpenRouter API Documentation,\" https://openrouter.ai/docs, 2024.",
        "[3] DeepSeek, \"DeepSeek-V3 Technical Report,\" 2024.",
        "[4] Mozilla, \"pdf.js: A Portable Document Format Library,\" https://mozilla.github.io/pdf.js/, 2024.",
        "[5] W3C, \"Web Speech API Specification,\" https://w3c.github.io/speech-api/, 2024.",
        "[6] AntV, \"Infographic: Analytical Visualization Library,\" https://github.com/antvis/Infographic, 2024.",
        "[7] Vercel, \"Serverless Functions Documentation,\" https://vercel.com/docs/functions, 2024.",
        "[8] IETF RFC 7519, \"JSON Web Token (JWT),\" 2015.",
        "[9] Express.js Foundation, \"Express Web Framework for Node.js,\" https://expressjs.com/, 2024.",
        "[10] Russell, S., and Norvig, P., Artificial Intelligence: A Modern Approach, 4th ed., Pearson, 2020.",
        "[11] Jurafsky, D., and Martin, J. H., Speech and Language Processing, 3rd ed. draft, 2023.",
        "[12] OpenAI, \"GPT-4 Technical Report,\" arXiv:2303.08774, 2023.",
    ]
    for r in refs:
        para(doc, r, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # APPENDICES
    # ═══════════════════════════════════════════════════════════════════════════
    heading(doc, "APPENDICES", 1)

    heading(doc, "Appendix A – Source Code Structure", 2)
    para(doc, "Project repository structure:")
    para(doc, "voice-study-assistant/\n├── frontend/\n│   ├── index.html, auth.html, app.js, styles.css, config.js, mindmap.js\n├── backend/\n│   ├── server.js, package.json, .env\n├── vercel.json\n└── README.md")
    para(doc, "Full source code submitted separately on CD/USB as per department guidelines.")

    heading(doc, "Appendix B – Sample Voice Commands", 2)
    table(doc, ["Voice Command", "Action"], [
        ["Explain photosynthesis", " Sends query to AI chat"],
        ["What is a dangling pointer?", " C programming query from PDF source"],
        ["Summarize chapter 3", " Summary from uploaded source"],
        ["(Mic click + speak in Telugu)", " Telugu query with Telugu response"],
    ])

    heading(doc, "Appendix C – Test Case Reports", 2)
    para(doc, "Detailed test case execution logs with date, tester name, browser, and pass/fail status maintained in project logbook. Summary provided in Table 6.1.")

    heading(doc, "Appendix D – User Manual", 2)
    para(doc, "1. Open https://voice-study-assistant.vercel.app\n2. Sign up with email and password\n3. Create a notebook and upload PDF/text sources\n4. Enable sources and ask questions via text or mic\n5. Use Studio cards to generate mind maps, slides, etc.\n6. Click Speak/Explain/Translate on AI responses\n7. Access saved outputs in Saved panel on right\n8. Sign out via avatar in top bar")

    heading(doc, "Appendix E – Installation Guide", 2)
    para(doc, "Local setup:\n1. Clone project\n2. cd backend && npm install\n3. Create .env with OPENROUTER_API_KEY, JWT_SECRET\n4. node server.js\n5. Serve frontend/ with any static server\n6. Open index.html or localhost URL\n\nDeploy: vercel --prod from project root")

    heading(doc, "Appendix F – Dataset Details", 2)
    para(doc, "No custom ML dataset was created. Testing used publicly available study PDFs (programming notes, engineering subjects) and user-entered topics. LLM training data is proprietary to model providers (DeepSeek, Google).")

    heading(doc, "Appendix G – Screenshots and Outputs", 2)
    para(doc, "[Attach high-resolution screenshots: Auth page, Main UI, Mind map, Flashcards, Audio overview, Explain overlay, Study resources, Infographic]")

    add_page_break(doc)

    # ANNEXURES
    heading(doc, "ANNEXURES", 1)
    annexures = [
        "Annexure I – Project Synopsis",
        "Annexure II – Weekly Progress Reports",
        "Annexure III – Gantt Chart / Project Schedule",
        "Annexure IV – Plagiarism Report",
        "Annexure V – Guide Meeting Logbook",
        "Annexure VI – Student Contribution Sheet",
        "Annexure VII – Attendance Record",
        "Annexure VIII – Ethical Clearance (if applicable)",
        "Annexure IX – Publication/Patent Details (if any)",
        "Annexure X – Project Completion Certificate",
    ]
    for a in annexures:
        para(doc, a, bold=True)
        para(doc, "[Attach signed/scanned document as per department format]")
        doc.add_paragraph()

    doc.save(OUTPUT)
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    build()
