"""Generate final B.Tech Major Project Report — no placeholders."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(BASE, "report_assets")
OUTPUT = os.path.join(BASE, "B.Tech_Major_Project_Report_FINAL.docx")

# ─── Student & institute details ─────────────────────────────────────────────
STUDENT_NAME    = "Maram Sai Shashank Raj"
ROLL_NUMBER     = "B24AI185"
UNIVERSITY      = "Jawaharlal Nehru Technological University, Hyderabad"
COLLEGE         = "Kakatiya Institute of Technology and Sciences, Warangal"
DEPARTMENT      = "Department of Computer Science and Engineering (AI & ML)"
BRANCH          = "COMPUTER SCIENCE AND ENGINEERING (AI & ML)"
GUIDE_NAME      = "Prof. Maram Balajee"
GUIDE_DESIG     = "Professor, Department of Computer Science and Engineering, SR University, Warangal"
ACADEMIC_YEAR   = "2025–26"
SUBMISSION_DATE = "June 2026"
PLACE           = "Warangal"
LIVE_URL        = "https://voice-study-assistant.vercel.app/"


def set_margins(doc):
    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.17)
        s.right_margin = Cm(2.54)


def pb(doc):
    doc.add_page_break()


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Times New Roman"
    return p


def p(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, sa=6):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(sa)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    return para


def bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    para.paragraph_format.line_spacing = 1.5
    for r in para.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for para in t.rows[0].cells[i].paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for para in t.rows[ri + 1].cells[ci].paragraphs:
                for r in para.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()


def center(doc, lines, size=12, bold=False):
    for line in lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def fig(doc, caption, filename, width=5.8):
    path = os.path.join(ASSETS, filename)
    if os.path.isfile(path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.font.name = "Times New Roman"
        cr.font.size = Pt(11)
        cr.italic = True
        doc.add_paragraph()


def make_architecture_diagram():
    path = os.path.join(ASSETS, "fig_architecture.png")
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    boxes = [
        (0.3, 2.8, 2.4, 1.2, "Client Browser\n(HTML/CSS/JS)\nWeb Speech API\nlocalStorage", "#4F46E5"),
        (3.8, 2.8, 2.4, 1.2, "Vercel Serverless\nNode.js + Express\nJWT Auth / SSE", "#7C3AED"),
        (7.3, 2.8, 2.4, 1.2, "OpenRouter API\nDeepSeek v3.2\nGemini TTS", "#059669"),
        (3.8, 0.8, 2.4, 1.0, "Static Frontend\nindex.html / auth.html", "#6366F1"),
    ]
    for x, y, w, h, label, color in boxes:
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05", facecolor=color, edgecolor="white", alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha="center", va="center", color="white", fontsize=9, fontweight="bold")
    ax.annotate("", xy=(3.7, 3.4), xytext=(2.8, 3.4), arrowprops=dict(arrowstyle="->", color="#333", lw=2))
    ax.annotate("", xy=(7.2, 3.4), xytext=(6.3, 3.4), arrowprops=dict(arrowstyle="->", color="#333", lw=2))
    ax.annotate("", xy=(5.0, 2.7), xytext=(5.0, 1.9), arrowprops=dict(arrowstyle="->", color="#333", lw=2))
    ax.set_title("Fig. 3.1 — System Architecture Overview", fontsize=12, fontweight="bold", pad=12)
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def make_dfd0():
    path = os.path.join(ASSETS, "fig_dfd0.png")
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis("off")
    ax.add_patch(mpatches.FancyBboxPatch((3.5, 1.5), 3, 1.2, boxstyle="round,pad=0.05", facecolor="#7C3AED", edgecolor="white"))
    ax.text(5, 2.1, "0\nVoice-Controlled\nAI Study Assistant", ha="center", va="center", color="white", fontsize=10, fontweight="bold")
    ax.text(0.5, 2.1, "Student /\nMentor", ha="center", va="center", fontsize=10, bbox=dict(boxstyle="circle", facecolor="#E0E7FF"))
    ax.text(9.5, 2.1, "OpenRouter\nLLM / TTS", ha="center", va="center", fontsize=10, bbox=dict(boxstyle="circle", facecolor="#D1FAE5"))
    ax.annotate("Queries, sources,\nvoice input", xy=(3.4, 2.1), xytext=(1.3, 2.1), arrowprops=dict(arrowstyle="<->", color="#333", lw=1.5))
    ax.annotate("AI requests /\nresponses", xy=(6.6, 2.1), xytext=(8.2, 2.1), arrowprops=dict(arrowstyle="<->", color="#333", lw=1.5))
    ax.set_title("Fig. 4.2 — Data Flow Diagram (Level 0)", fontsize=12, fontweight="bold")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()


def build():
    os.makedirs(ASSETS, exist_ok=True)
    make_architecture_diagram()
    make_dfd0()

    doc = Document()
    set_margins(doc)

    # COVER
    doc.add_paragraph()
    center(doc, [UNIVERSITY.upper()], 16, True)
    center(doc, [COLLEGE.upper()], 14, True)
    doc.add_paragraph()
    center(doc, ["A Major Project Report"], 14, True)
    center(doc, ["on"], 12)
    center(doc, ['"VOICE-CONTROLLED AI STUDY ASSISTANT"'], 16, True)
    center(doc, ["(An AI-Powered NotebookLM-Style Learning Platform)"], 12)
    doc.add_paragraph()
    center(doc, ["Submitted in partial fulfilment of the requirements"], 12)
    center(doc, ["for the award of the degree of"], 12)
    center(doc, ["BACHELOR OF TECHNOLOGY"], 14, True)
    center(doc, ["in"], 12)
    center(doc, [BRANCH], 14, True)
    doc.add_paragraph()
    center(doc, [f"Academic Year: {ACADEMIC_YEAR}"], 12)
    doc.add_paragraph()
    center(doc, ["Submitted by"], 12)
    center(doc, [STUDENT_NAME], 12, True)
    center(doc, [f"Roll No.: {ROLL_NUMBER}"], 12)
    doc.add_paragraph()
    center(doc, ["Under the guidance of"], 12)
    center(doc, [GUIDE_NAME], 12, True)
    center(doc, [GUIDE_DESIG], 12)
    center(doc, [DEPARTMENT], 12)
    center(doc, [COLLEGE], 12)
    doc.add_paragraph()
    center(doc, [SUBMISSION_DATE], 12)
    pb(doc)

    # BONAFIDE
    h(doc, "BONAFIDE CERTIFICATE", 1)
    p(doc, f'This is to certify that the Major Project Report entitled "VOICE-CONTROLLED AI STUDY ASSISTANT" submitted by {STUDENT_NAME} (Roll No.: {ROLL_NUMBER}) in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in {BRANCH.title()} is a bonafide record of work carried out by him under my supervision and guidance.')
    p(doc, "The matter embodied in this report has not been submitted elsewhere for the award of any other degree or diploma.")
    doc.add_paragraph()
    p(doc, f"Date: {SUBMISSION_DATE}", align=WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, f"Place: {PLACE}\n\nSignature of the Guide\n{GUIDE_NAME}\n{GUIDE_DESIG}", align=WD_ALIGN_PARAGRAPH.LEFT)
    pb(doc)

    # APPROVAL
    h(doc, "PROJECT APPROVAL SHEET", 1)
    p(doc, f'The Major Project Report entitled "VOICE-CONTROLLED AI STUDY ASSISTANT" submitted by {STUDENT_NAME} has been examined and approved for submission.')
    tbl(doc, ["Name", "Roll Number", "Signature"], [[STUDENT_NAME, ROLL_NUMBER, ""]])
    p(doc, f"Project Guide: {GUIDE_NAME}          Date: {SUBMISSION_DATE}")
    p(doc, "Internal Examiner: _____________________          Date: __________")
    p(doc, "External Examiner: _____________________          Date: __________")
    pb(doc)

    # DECLARATION
    h(doc, "DECLARATION", 1)
    p(doc, f'I, {STUDENT_NAME}, hereby declare that the Major Project Report entitled "VOICE-CONTROLLED AI STUDY ASSISTANT" submitted to {UNIVERSITY} through {COLLEGE} in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in {BRANCH.title()} is an authentic record of my own work carried out under the supervision of {GUIDE_NAME}.')
    p(doc, "The matter embodied in this report has not been submitted by me for the award of any other degree or diploma. Wherever other works have been referred to, due acknowledgement has been made.")
    doc.add_paragraph()
    p(doc, f"Date: {SUBMISSION_DATE}")
    p(doc, f"Place: {PLACE}")
    doc.add_paragraph()
    p(doc, "Signature of the Student: _________________________")
    p(doc, STUDENT_NAME)
    pb(doc)

    # ACKNOWLEDGEMENT
    h(doc, "ACKNOWLEDGEMENT", 1)
    p(doc, f"I express my sincere gratitude to my project guide, {GUIDE_NAME}, {GUIDE_DESIG}, for invaluable guidance, continuous encouragement, and constructive feedback throughout the development of this project.")
    p(doc, f"I thank the faculty members of {DEPARTMENT}, {COLLEGE}, for providing academic support and laboratory facilities.")
    p(doc, "I acknowledge OpenRouter, DeepSeek, Google Gemini, Mozilla pdf.js, AntV, and the open-source community for tools and APIs that made this project feasible.")
    p(doc, "Finally, I thank my parents and friends for their moral support and motivation.")
    pb(doc)

    # ABSTRACT
    h(doc, "ABSTRACT", 1)
    p(doc, "Traditional study methods often require students to manually organize notes, search across multiple documents, and switch between reading, revision, and assessment tools. Generic AI chatbots provide answers but lack source grounding, persistent notebooks, and integrated study workflows inspired by modern research assistants such as Google NotebookLM.")
    p(doc, f"This project presents a Voice-Controlled AI Study Assistant — a full-stack web application deployed at {LIVE_URL} that enables students to upload study sources (PDF/text), interact through text and voice, and generate rich learning artifacts including mind maps, slide decks, flashcards, FAQs, briefing documents, study outlines, data tables, study resource roadmaps, and infographics. The system uses DeepSeek v3.2 via OpenRouter for conversational AI, JSON-based chat responses on Vercel serverless, Google Gemini TTS for multilingual speech output, and the Web Speech API for speech-to-text input.")
    p(doc, "The application follows a three-panel NotebookLM-inspired interface: Sources, Chat, and Studio. User authentication is implemented using stateless JWT-based identity derivation, with per-user data isolation in browser localStorage. Experimental evaluation demonstrates successful deployment, multilingual query handling, and effective generation of study materials from uploaded sources.")
    p(doc, "Keywords: Artificial Intelligence, Voice-Controlled Systems, Natural Language Processing, Study Assistant, OpenRouter, Speech-to-Text, Text-to-Speech, NotebookLM, Web Application, Vercel.")
    pb(doc)

    # TOC
    h(doc, "TABLE OF CONTENTS", 1)
    for line in [
        "CHAPTER 1 – INTRODUCTION",
        "CHAPTER 2 – LITERATURE SURVEY",
        "CHAPTER 3 – SYSTEM ANALYSIS",
        "CHAPTER 4 – DESIGN METHODOLOGY",
        "CHAPTER 5 – IMPLEMENTATION",
        "CHAPTER 6 – RESULTS AND DISCUSSION",
        "CHAPTER 7 – CONCLUSION AND FUTURE SCOPE",
        "REFERENCES",
        "APPENDICES",
    ]:
        p(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT)
    pb(doc)

    # LIST OF FIGURES
    h(doc, "LIST OF FIGURES", 1)
    for line in [
        "Fig. 3.1  System Architecture Overview",
        "Fig. 4.2  Data Flow Diagram – Level 0",
        "Fig. 5.1  Authentication Page",
        "Fig. 5.2  Main Three-Panel Interface",
        "Fig. 5.3  Chat Interface with Saved Studio Outputs",
        "Fig. 5.4  Interactive Mind Map Output",
        "Fig. 6.1  Chat Response with Action Buttons and Code Explanation",
    ]:
        p(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT)
    pb(doc)

    # ABBREVIATIONS
    h(doc, "LIST OF ABBREVIATIONS AND ACRONYMS", 1)
    tbl(doc, ["Abbreviation", "Expansion"], [
        ["AI", "Artificial Intelligence"], ["API", "Application Programming Interface"],
        ["CSS", "Cascading Style Sheets"], ["DFD", "Data Flow Diagram"],
        ["HTML", "HyperText Markup Language"], ["JWT", "JSON Web Token"],
        ["LLM", "Large Language Model"], ["NLP", "Natural Language Processing"],
        ["PDF", "Portable Document Format"], ["SSE", "Server-Sent Events"],
        ["STT", "Speech-to-Text"], ["SVG", "Scalable Vector Graphics"],
        ["TTS", "Text-to-Speech"], ["UI", "User Interface"],
    ])
    pb(doc)

    # ─── CHAPTER 1 ───────────────────────────────────────────────────────────
    h(doc, "CHAPTER 1", 1)
    h(doc, "INTRODUCTION", 1)
    h(doc, "1.1 Introduction to Artificial Intelligence", 2)
    p(doc, "Artificial Intelligence (AI) refers to the simulation of human intelligence in machines programmed to think, learn, and solve problems. In education, AI enables personalized tutoring, automated content summarization, intelligent question answering, and adaptive learning pathways.")
    h(doc, "1.2 Voice-Controlled Systems Overview", 2)
    p(doc, "Voice-controlled systems allow users to interact with software through spoken commands. The Web Speech API enables browser-native speech recognition and synthesis, supporting hands-free learning during revision sessions.")
    h(doc, "1.3 Need for AI-Based Study Assistants", 2)
    p(doc, "Students manage multiple subjects, PDFs, lecture notes, and revision schedules simultaneously. An integrated AI study assistant combines source management, grounded chat, voice interaction, and automated study artifact generation in one workspace.")
    h(doc, "1.4 Problem Statement", 2)
    p(doc, "The problem addressed is the lack of an affordable, deployable AI study platform that grounds responses in user-uploaded sources, supports voice and multilingual output, generates diverse study materials, persists notebooks across sessions, and isolates data between users.")
    h(doc, "1.5 Motivation of the Project", 2)
    p(doc, "Commercial platforms such as Google NotebookLM demonstrate the value of source-grounded AI for learning. This project builds a college-level, deployable alternative using OpenRouter, modern web technologies, and a NotebookLM-inspired user experience.")
    h(doc, "1.6 Objectives of the Project", 2)
    for o in [
        "Design and implement a web-based AI study assistant with NotebookLM-style three-panel layout.",
        "Integrate voice input (STT) and voice output (TTS) for hands-free learning.",
        "Enable PDF/text source upload and source-grounded AI chat.",
        "Implement Studio features: mind maps, slides, flashcards, FAQ, briefing, outline, data tables, study resources, and infographics.",
        "Provide audio overview podcasts with interactive join-conversation and continue-in-podcast modes.",
        "Implement user authentication with per-user data isolation.",
        f"Deploy the application on Vercel at {LIVE_URL}.",
    ]:
        bullet(doc, o)
    h(doc, "1.7 Scope of the Project", 2)
    p(doc, "Scope includes frontend UI, backend API proxy, OpenRouter LLM integration, authentication, local persistence, and cloud deployment. Custom LLM training and native mobile applications are out of scope.")
    h(doc, "1.8 Proposed Solution", 2)
    p(doc, "The proposed system is a full-stack web application with Vanilla JavaScript frontend, Node.js/Express backend on Vercel, OpenRouter for LLM/TTS, JWT stateless authentication, and localStorage-based notebook persistence scoped per user.")
    h(doc, "1.9 Applications of the System", 2)
    for a in ["Engineering exam preparation", "Competitive exam revision", "Research paper summarization", "Multilingual language learning", "Classroom AI demonstration", "Remote mentoring via shared link"]:
        bullet(doc, a)
    h(doc, "1.10 Advantages of the Proposed System", 2)
    for a in ["Unified workspace for sources, chat, and study content", "Voice and text multimodal interaction", "Low cost via Vercel and OpenRouter pay-per-use", "Browser-based — no installation", "Rich studio outputs beyond plain chat", "Per-user data isolation"]:
        bullet(doc, a)
    h(doc, "1.11 Organization of the Report", 2)
    p(doc, "Chapter 2 presents literature survey. Chapter 3 covers system analysis. Chapter 4 describes design. Chapter 5 explains implementation. Chapter 6 discusses results. Chapter 7 concludes with future scope.")
    pb(doc)

    # ─── CHAPTER 2 (condensed key sections) ──────────────────────────────────
    h(doc, "CHAPTER 2", 1)
    h(doc, "LITERATURE SURVEY", 1)
    h(doc, "2.1 Introduction", 2)
    p(doc, "This chapter reviews voice assistants, AI in education, speech technologies, NLP approaches, and comparative systems.")
    h(doc, "2.7 Comparative Analysis of Existing Systems", 2)
    tbl(doc, ["System", "Source Grounding", "Voice", "Study Artifacts", "Custom Deploy"], [
        ["Google NotebookLM", "Yes", "Limited", "Audio, slides", "No"],
        ["ChatGPT", "Via uploads", "Yes (Plus)", "Generic", "No"],
        ["Proposed System", "Yes", "Yes (STT+TTS)", "Mind map, flashcards, podcast", "Yes (Vercel)"],
    ])
    h(doc, "2.9 Summary", 2)
    p(doc, "Literature review confirms demand for integrated AI study platforms. The proposed system fills gaps by combining NotebookLM-inspired UX, voice control, rich studio outputs, and practical deployment.")
    pb(doc)

    # ─── CHAPTER 3 ─────────────────────────────────────────────────────────
    h(doc, "CHAPTER 3", 1)
    h(doc, "SYSTEM ANALYSIS", 1)
    h(doc, "3.3 Proposed System", 2)
    p(doc, f"The proposed Voice-Controlled AI Study Assistant is deployed at {LIVE_URL} with authentication, multi-notebook management, source upload, AI chat, Studio generation, audio overview, explain/translate/speak actions, and saved outputs.")
    h(doc, "3.5.1 Hardware Requirements", 3)
    tbl(doc, ["Component", "Specification"], [
        ["Client PC/Laptop", "Intel i3 or equivalent, 4 GB RAM"],
        ["Microphone", "Built-in or external"],
        ["Network", "Broadband internet"],
    ])
    h(doc, "3.5.2 Software Requirements", 3)
    tbl(doc, ["Software", "Purpose"], [
        ["Chrome / Edge 90+", "Web browser"],
        ["Node.js 18+", "Local development"],
        ["OpenRouter API Key", "LLM and TTS access"],
        ["Vercel", "Cloud deployment"],
    ])
    h(doc, "3.8 System Architecture Overview", 2)
    p(doc, "The system follows three tiers: Presentation (HTML/CSS/JS + Web Speech API), Application (Express.js on Vercel serverless), and External Services (OpenRouter DeepSeek/Gemini). Data persists in browser localStorage keyed by user ID.")
    fig(doc, "Fig. 3.1 — System Architecture Overview", "fig_architecture.png", 6.0)
    pb(doc)

    # ─── CHAPTER 4 ─────────────────────────────────────────────────────────
    h(doc, "CHAPTER 4", 1)
    h(doc, "DESIGN METHODOLOGY", 1)
    h(doc, "4.2 Architecture of the Proposed System", 2)
    p(doc, "Monorepo: frontend/ (static) and backend/ (server.js). Vercel routes /api/* to serverless backend. config.js switches API_BASE between localhost and same-origin production.")
    h(doc, "4.3 Module Description", 2)
    for title, body in [
        ("4.3.1 Voice Input Module", "Web Speech API captures microphone input; final transcript sent to chat pipeline."),
        ("4.3.2 Speech-to-Text Module", "Browser STT with regex-based language detection for Telugu, Hindi, Tamil, and other scripts."),
        ("4.3.3 NLP Processing Module", "LLM prompt orchestration injects sources, notes, and chat history into OpenRouter requests."),
        ("4.3.4 AI Response Generation Module", "POST /api/chat returns JSON on Vercel; SSE streaming used locally."),
        ("4.3.5 Data Storage Module", "localStorage keys: sa_auth, sa_notebooks_v2_<uid> for notebooks, sources, chats, studioOutputs."),
        ("4.3.6 User Interface Module", "Three-panel layout: Sources | Chat | Studio with premium auth.html login page."),
    ]:
        h(doc, title, 3)
        p(doc, body)
    h(doc, "4.4 Data Flow Diagram", 2)
    fig(doc, "Fig. 4.2 — Data Flow Diagram (Level 0)", "fig_dfd0.png", 5.5)
    h(doc, "4.5 Use Case Diagram", 2)
    p(doc, "Actors: Student, Mentor, OpenRouter API. Use cases: Sign Up, Sign In, Upload Source, Chat (Text/Voice), Generate Mind Map, Slides, Flashcards, Audio Overview, Explain, Translate, Save Output, Sign Out.")
    h(doc, "4.11 Security and Privacy Considerations", 2)
    for s in ["API key stored only in server environment variables", "JWT HS256 with 365-day expiry", "Per-user localStorage isolation", "HTTPS via Vercel"]:
        bullet(doc, s)
    pb(doc)

    # ─── CHAPTER 5 ─────────────────────────────────────────────────────────
    h(doc, "CHAPTER 5", 1)
    h(doc, "IMPLEMENTATION", 1)
    h(doc, "5.2 Development Environment", 2)
    p(doc, f"Development on Windows with VS Code/Cursor. Local backend on port 3000. Production at {LIVE_URL} via Vercel CLI.")
    h(doc, "5.3 Tools and Technologies Used", 2)
    tbl(doc, ["Category", "Technology", "Purpose"], [
        ["Frontend", "HTML5, CSS3, JavaScript", "UI and client logic"],
        ["Backend", "Node.js, Express.js", "REST API"],
        ["AI", "DeepSeek v3.2 via OpenRouter", "Chat and generation"],
        ["TTS", "Gemini 3.1 Flash TTS", "Multilingual speech"],
        ["PDF", "pdf.js", "Text extraction"],
        ["Mind Map", "SVG + Reingold-Tilford", "Interactive maps"],
        ["Deployment", "Vercel", "Static + serverless"],
    ])
    h(doc, "5.8 User Authentication Process", 2)
    p(doc, "Stateless auth: uid = HMAC(email + password). JWT stored in localStorage. Mentor login via username ss6156. Auth guard redirects unauthenticated users to auth.html.")
    h(doc, "5.10 Screenshots of Implementation", 2)
    fig(doc, "Fig. 5.1 — Authentication Page with Sign In / Sign Up and Mentor Portal", "fig_auth.png")
    fig(doc, "Fig. 5.2 — Main Three-Panel Interface (Sources | Chat | Studio)", "fig_main_ui.png")
    fig(doc, "Fig. 5.3 — Chat Interface with Saved Studio Outputs Panel", "fig_chat_saved.png")
    fig(doc, "Fig. 5.4 — Interactive Mind Map Generated from PDF Source", "fig_mindmap.png")
    pb(doc)

    # ─── CHAPTER 6 ─────────────────────────────────────────────────────────
    h(doc, "CHAPTER 6", 1)
    h(doc, "RESULTS AND DISCUSSION", 1)
    h(doc, "6.2 Test Cases", 2)
    tbl(doc, ["TC ID", "Test Case", "Result"], [
        ["TC-01", "User sign-up with email", "Pass"],
        ["TC-02", "User sign-in — notebooks loaded", "Pass"],
        ["TC-03", "Upload PDF source", "Pass"],
        ["TC-04", "Chat query on source (pointer to pointer)", "Pass"],
        ["TC-05", "Voice input query", "Pass"],
        ["TC-06", "Generate mind map", "Pass"],
        ["TC-07", "Generate flashcards, slides, audio", "Pass"],
        ["TC-08", "Explain / Translate / Speak buttons", "Pass"],
        ["TC-09", "Cross-device Vercel access", "Pass"],
        ["TC-10", "Mentor login", "Pass"],
    ])
    h(doc, "6.5 Response Time Analysis", 2)
    tbl(doc, ["Operation", "Average Time"], [
        ["Sign-in API", "1–6 seconds"],
        ["Full chat response", "15–25 seconds (Vercel)"],
        ["Mind map generation", "15–30 seconds"],
        ["PDF extraction (20 pages)", "3–8 seconds"],
    ])
    h(doc, "6.8 Discussion of Results", 2)
    p(doc, f"The system successfully demonstrates NotebookLM-style learning at {LIVE_URL}. Uploading 'two pointers complete base' PDF and querying 'define pointer to pointer' produced accurate C/C++ code explanations with saved mind maps, flashcards, slides, and audio outputs visible in the Saved panel.")
    h(doc, "6.9 Limitations", 2)
    for l in ["Data in localStorage — not synced across devices", "No vector RAG for very large PDFs", "Vercel cold starts cause initial delay", "STT accuracy varies by browser"]:
        bullet(doc, l)
    pb(doc)

    # ─── CHAPTER 7 ─────────────────────────────────────────────────────────
    h(doc, "CHAPTER 7", 1)
    h(doc, "CONCLUSION AND FUTURE SCOPE", 1)
    h(doc, "7.1 Conclusion", 2)
    p(doc, f"The Voice-Controlled AI Study Assistant successfully delivers an integrated AI-powered learning platform at {LIVE_URL}, combining source-grounded chat, voice interaction, studio artifacts, audio podcasts, and user authentication.")
    h(doc, "7.2 Achievements", 2)
    for a in ["Full-stack NotebookLM-style UI implemented", "DeepSeek v3.2 integrated via OpenRouter", "10+ studio features built", "Voice input and TTS output working", "Deployed on Vercel with authentication", "Complete project documentation prepared"]:
        bullet(doc, a)
    h(doc, "7.3 Future Enhancements", 2)
    for f in ["Cloud database for cross-device sync", "Vector RAG for large documents", "Mobile PWA app", "Collaborative group notebooks"]:
        bullet(doc, f)
    h(doc, "7.5 Final Remarks", 2)
    p(doc, "This project demonstrates that modern LLM APIs, browser speech technologies, and serverless deployment can build sophisticated educational tools suitable for academic demonstration and real-world learning.")
    pb(doc)

    # REFERENCES
    h(doc, "REFERENCES", 1)
    for r in [
        "[1] Google, NotebookLM: An AI-powered research and writing assistant, Google Labs, 2023–2024.",
        "[2] OpenRouter, OpenRouter API Documentation, https://openrouter.ai/docs, 2024.",
        "[3] DeepSeek, DeepSeek-V3 Technical Report, 2024.",
        "[4] Mozilla, pdf.js Documentation, https://mozilla.github.io/pdf.js/, 2024.",
        "[5] W3C, Web Speech API Specification, 2024.",
        "[6] Vercel, Serverless Functions Documentation, https://vercel.com/docs/functions, 2024.",
        "[7] IETF RFC 7519, JSON Web Token (JWT), 2015.",
        "[8] Russell, S., and Norvig, P., Artificial Intelligence: A Modern Approach, 4th ed., Pearson, 2020.",
    ]:
        p(doc, r, align=WD_ALIGN_PARAGRAPH.LEFT)
    pb(doc)

    # APPENDICES
    h(doc, "APPENDICES", 1)
    h(doc, "Appendix A – Source Code Structure", 2)
    p(doc, "voice-study-assistant/\n├── frontend/ (index.html, auth.html, app.js, styles.css, mindmap.js)\n├── backend/ (server.js, package.json)\n├── vercel.json\n└── README.md")
    h(doc, "Appendix D – User Manual", 2)
    p(doc, f"1. Open {LIVE_URL}\n2. Sign up with email and password\n3. Create a notebook and upload PDF/text sources\n4. Ask questions via text or microphone\n5. Use Studio cards to generate mind maps, slides, flashcards, etc.\n6. Click Speak, Explain, or Translate on AI responses\n7. Access saved outputs in the Saved panel\n8. Sign out via avatar in top bar")
    h(doc, "Appendix E – Installation Guide", 2)
    p(doc, "Local: cd backend && npm install && node server.js. Serve frontend/ statically.\nDeploy: vercel --prod from project root with OPENROUTER_API_KEY set.")

    h(doc, "ANNEXURES", 1)
    p(doc, "Annexures (Project Synopsis, Weekly Progress Reports, Gantt Chart, Plagiarism Report, Guide Logbook, and Completion Certificate) are submitted separately to the department as per institute norms.")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
