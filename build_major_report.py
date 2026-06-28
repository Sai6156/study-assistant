"""
Build B.Tech Major Project Report — Voice-Controlled AI Study Assistant
Formatting matched to SmartAgri reference document (KITS Warangal).
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from report_diagrams import generate_all
from report_toc import TOC_ENTRIES
from chapter4_design import build_chapter4 as build_chapter4_detailed

try:
    from capture_screenshots import capture as capture_screenshots
except ImportError:
    capture_screenshots = None

BASE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(BASE, "report_assets")
OUTPUT = os.path.join(BASE, "B.Tech_Major_Project_Report_Voice_Controlled_AI_Study_Assistant.docx")
OUTPUT_ALT = os.path.join(BASE, "B.Tech_Major_Project_Report_FINAL.docx")

# ─── Student & institute details ───────────────────────────────────────────────
STUDENT_NAME = "Maram Sai Shashank Raj"
ROLL_NUMBER = "B24AI185"
COLLEGE = "Kakatiya Institute of Technology and Sciences, Warangal"
DEPARTMENT = "Department of Computer Science and Engineering (AI & ML)"
BRANCH = "COMPUTER SCIENCE AND ENGINEERING (AI & ML)"
GUIDE_NAME = "Prof. Maram Balajee"
GUIDE_DESIG = "Professor, Department of Computer Science and Engineering, SR University, Warangal"
ACADEMIC_YEAR = "2025–26"
SUBMISSION_DATE = "June 2026"
PLACE = "Warangal"
LIVE_URL = "https://voice-study-assistant.vercel.app/"
PROJECT_TITLE = "VOICE-CONTROLLED AI STUDY ASSISTANT"
PROJECT_SUBTITLE = "(An AI-Powered NotebookLM-Style Learning Platform)"

FONT = "Times New Roman"
BODY_SIZE = Pt(12)
CAPTION_SIZE = Pt(10)
LINE_SPACING = 1.5

# ─── Low-level OOXML helpers ─────────────────────────────────────────────────

def _set_run_font(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def _page_field(paragraph):
    for part, text in [("begin", None), (None, " PAGE  \\* MERGEFORMAT "), ("separate", None), (None, "1"), ("end", None)]:
        run = paragraph.add_run()
        _set_run_font(run, Pt(11))
        if text is None:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), part)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr) if "MERGEFORMAT" in text else run._element.append(OxmlElement("w:t"))


def _add_page_number_footer(section, fmt="decimal", start=None):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run()
    _set_run_font(run, Pt(11))
    # PAGE field
    for step in ("begin", "instr", "sep", "end"):
        r = p.add_run()
        _set_run_font(r, Pt(11))
        if step == "begin":
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r._r.append(fc)
        elif step == "instr":
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
            it.text = " PAGE  \\* MERGEFORMAT "; r._r.append(it)
        elif step == "sep":
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate"); r._r.append(fc)
            r.add_text("1")
        else:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end"); r._r.append(fc)

    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    elif qn("w:start") in pgNumType.attrib:
        del pgNumType.attrib[qn("w:start")]


def _set_margins(section):
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)


def new_section(doc, fmt=None, start=None):
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec = doc.sections[-1]
    _set_margins(sec)
    if fmt:
        _add_page_number_footer(sec, fmt, start=start)
    return sec


# ─── Document helpers ────────────────────────────────────────────────────────

def pb(doc):
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        _set_run_font(r, Pt(14), bold=True)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        _set_run_font(r, Pt(13), bold=True)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        _set_run_font(r, Pt(12), bold=True)
    return p


def para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=BODY_SIZE):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, size, bold=bold)
    return p


def center(doc, lines, size=12, bold=False):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(line)
        _set_run_font(run, Pt(size), bold=bold)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.line_spacing = LINE_SPACING
    for r in p.runs:
        _set_run_font(r, BODY_SIZE)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _set_run_font(run, CAPTION_SIZE, bold=True)


def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, CAPTION_SIZE, bold=True)


def tbl(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hdr in enumerate(headers):
        t.rows[0].cells[i].text = hdr
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                _set_run_font(r, Pt(11), bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    _set_run_font(r, Pt(11))
    doc.add_paragraph()
    return t


def fig(doc, caption_text, filename, width=5.8):
    path = os.path.join(ASSETS, filename)
    if not os.path.isfile(path):
        # try without path prefix
        alt = filename
        if not os.path.isfile(alt):
            para(doc, f"[Figure placeholder: {caption_text}]", align=WD_ALIGN_PARAGRAPH.CENTER)
            caption(doc, caption_text)
            return
        path = alt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    caption(doc, caption_text)


def toc_entry(doc, title, page, indent_level=0):
    if not title:
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.left_indent = Pt(12 * indent_level)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.35), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    run = p.add_run(f"{title}\t{page}")
    _set_run_font(run, BODY_SIZE)


# ─── Preliminary pages ─────────────────────────────────────────────────────────

def build_cover_and_bonafide(doc):
    """Compact cover + bonafide on one page (merged opening pages)."""
    center(doc, [COLLEGE.upper()], 16, True)
    center(doc, ["A Major Project Report"], 14, True)
    center(doc, ["on"], 12)
    center(doc, [f'"{PROJECT_TITLE}"'], 15, True)
    center(doc, [PROJECT_SUBTITLE], 11)
    center(doc, ["Submitted in partial fulfilment of the requirements"], 11)
    center(doc, ["for the award of the degree of"], 11)
    center(doc, ["BACHELOR OF TECHNOLOGY"], 13, True)
    center(doc, ["in"], 11)
    center(doc, [BRANCH], 13, True)
    center(doc, [f"Academic Year: {ACADEMIC_YEAR}"], 11)
    center(doc, ["Submitted by"], 11)
    center(doc, [STUDENT_NAME], 12, True)
    center(doc, [f"Roll No.: {ROLL_NUMBER}"], 11)
    center(doc, ["Under the guidance of"], 11)
    center(doc, [GUIDE_NAME], 11, True)
    center(doc, [GUIDE_DESIG], 10)
    center(doc, [DEPARTMENT], 10)
    center(doc, [SUBMISSION_DATE], 11)
    doc.add_paragraph()
    h1(doc, "BONAFIDE CERTIFICATE")
    para(doc, f'This is to certify that the Major Project Report entitled "{PROJECT_TITLE}" submitted by {STUDENT_NAME} (Roll No.: {ROLL_NUMBER}) in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering (AI & ML) is a bonafide record of work carried out by him under my supervision and guidance.')
    para(doc, "The matter embodied in this report has not been submitted elsewhere for the award of any other degree or diploma.")
    para(doc, f"Date: {SUBMISSION_DATE}")
    para(doc, f"Place: {PLACE}          Signature of the Guide          {GUIDE_NAME}, {GUIDE_DESIG}")
    pb(doc)


def build_preliminary(doc):
    h1(doc, "PROJECT APPROVAL SHEET")
    para(doc, f'The Major Project Report entitled "{PROJECT_TITLE}" submitted by {STUDENT_NAME} has been examined and approved for submission.')
    tbl(doc, ["Name", "Roll Number", "Signature"], [[STUDENT_NAME, ROLL_NUMBER, ""]])
    para(doc, f"Project Guide: {GUIDE_NAME}          Date: {SUBMISSION_DATE}")
    para(doc, "Internal Examiner: _____________________          Date: __________")
    para(doc, "External Examiner: _____________________          Date: __________")
    pb(doc)

    h1(doc, "DECLARATION")
    para(doc, f'I, {STUDENT_NAME}, hereby declare that the Major Project Report entitled "{PROJECT_TITLE}" submitted to {COLLEGE} in partial fulfilment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering (AI & ML) is an authentic record of my own work carried out under the supervision of {GUIDE_NAME}.')
    para(doc, "The matter embodied in this report has not been submitted by me for the award of any other degree or diploma. Wherever other works have been referred to, due acknowledgement has been made.")
    para(doc, f"Date: {SUBMISSION_DATE}")
    para(doc, f"Place: {PLACE}")
    para(doc, "Signature of the Student: _________________________")
    para(doc, STUDENT_NAME)
    pb(doc)

    h1(doc, "CERTIFICATE FROM SUPERVISOR")
    para(doc, f'This is to certify that the project work titled "{PROJECT_TITLE}" carried out by {STUDENT_NAME} (Roll No.: {ROLL_NUMBER}) under my guidance is a record of bonafide work and has been completed to my satisfaction. The report is recommended for submission to {COLLEGE} for evaluation.')
    para(doc, f"Date: {SUBMISSION_DATE}")
    para(doc, f"Place: {PLACE}")
    para(doc, f"Signature of Supervisor: _________________________{GUIDE_NAME}{GUIDE_DESIG}")
    pb(doc)

    h1(doc, "ACKNOWLEDGEMENT")
    para(doc, f"I express my sincere gratitude to my project guide, {GUIDE_NAME}, {GUIDE_DESIG}, for invaluable guidance, continuous encouragement, and constructive feedback throughout the development of this project.")
    para(doc, f"I thank the faculty members of {DEPARTMENT}, {COLLEGE}, for providing academic support and laboratory facilities.")
    para(doc, "I acknowledge OpenRouter, DeepSeek, Google Gemini, Mozilla pdf.js, AntV, and the open-source community for tools and APIs that made this project feasible.")
    para(doc, "Finally, I thank my parents and friends for their moral support and motivation.")
    pb(doc)

    h1(doc, "ABSTRACT")
    para(doc, "Traditional study methods often require students to manually organize notes, search across multiple documents, and switch between reading, revision, and assessment tools. Generic AI chatbots provide answers but lack source grounding, persistent notebooks, and integrated study workflows inspired by modern research assistants such as Google NotebookLM.")
    para(doc, f"This project presents a Voice-Controlled AI Study Assistant — a full-stack web application deployed at {LIVE_URL} that enables students to upload study sources (PDF/text), interact through text and voice, and generate rich learning artifacts including mind maps, slide decks, flashcards, FAQs, briefing documents, study outlines, data tables, study resource roadmaps, and infographics. The system uses DeepSeek v3.2 via OpenRouter for conversational AI, Server-Sent Events (SSE) for real-time streaming responses, Google Gemini TTS for multilingual speech output, and the Web Speech API for speech-to-text input.")
    para(doc, "The application follows a three-panel NotebookLM-inspired interface: Sources, Chat, and Studio. User authentication is implemented using stateless JWT-based identity derivation, with per-user data isolation in browser localStorage. Key features include audio overview podcasts, join-the-conversation mode, personalized audio explanations, and persistent saved studio outputs.")
    para(doc, "Keywords: Artificial Intelligence, Voice-Controlled Systems, Natural Language Processing, Study Assistant, OpenRouter, Speech-to-Text, Text-to-Speech, NotebookLM, Web Application, Vercel.")
    pb(doc)

    h1(doc, "TABLE OF CONTENTS")
    for title, page, indent in TOC_ENTRIES:
        toc_entry(doc, title, page, indent)
    pb(doc)

    h1(doc, "LIST OF FIGURES")
    for line in [
        "Fig. 3.1  System Architecture Overview",
        "Fig. 4.1  High-Level Architecture of Proposed System",
        "Fig. 4.2  Data Flow Diagram – Level 0",
        "Fig. 4.3  Data Flow Diagram – Level 1",
        "Fig. 4.4  Use Case Diagram",
        "Fig. 4.5  Sequence Diagram – Chat Interaction",
        "Fig. 4.6  Activity Diagram – User Study Session",
        "Fig. 4.7  Entity Relationship Diagram (Logical Data Model)",
        "Fig. 4.8  System Operation Flowchart",
        "Fig. 5.1  Authentication Page",
        "Fig. 5.2  Main Three-Panel Interface",
        "Fig. 5.3  Chat Interface with Saved Studio Outputs",
        "Fig. 5.4  Interactive Mind Map Output",
        "Fig. 6.1  Chat Response with Action Buttons",
        "Fig. 6.2  Study Resources Output",
    ]:
        para(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT)
    pb(doc)

    h1(doc, "LIST OF TABLES")
    for line in [
        "Table 2.1  Comparative Analysis of Existing Systems",
        "Table 3.1  Hardware Requirements",
        "Table 3.2  Software Requirements",
        "Table 3.3  Functional Requirements",
        "Table 3.4  Non-Functional Requirements",
        "Table 5.1  Technology Stack",
        "Table 6.1  Test Cases and Results",
        "Table 6.2  Performance Metrics",
    ]:
        para(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT)
    pb(doc)

    h1(doc, "LIST OF ABBREVIATIONS AND ACRONYMS")
    tbl(doc, ["Abbreviation", "Expansion"], [
        ["AI", "Artificial Intelligence"], ["API", "Application Programming Interface"],
        ["CSS", "Cascading Style Sheets"], ["DFD", "Data Flow Diagram"],
        ["DOM", "Document Object Model"], ["ER", "Entity Relationship"],
        ["HTML", "HyperText Markup Language"], ["JWT", "JSON Web Token"],
        ["LLM", "Large Language Model"], ["ML", "Machine Learning"],
        ["NLP", "Natural Language Processing"], ["PDF", "Portable Document Format"],
        ["REST", "Representational State Transfer"], ["SSE", "Server-Sent Events"],
        ["STT", "Speech-to-Text"], ["SVG", "Scalable Vector Graphics"],
        ["TTS", "Text-to-Speech"], ["UI", "User Interface"], ["UX", "User Experience"],
    ])
    pb(doc)


# ─── Chapters (abbreviated builder calls full content) ─────────────────────────

def build_chapter1(doc):
    h1(doc, "CHAPTER 1")
    h1(doc, "INTRODUCTION")
    h2(doc, "1.1 Introduction to Artificial Intelligence")
    para(doc, "Artificial Intelligence (AI) refers to the simulation of human intelligence in machines programmed to think, learn, and solve problems. In education, AI enables personalized tutoring, automated content summarization, intelligent question answering, and adaptive learning pathways. Large Language Models (LLMs) such as GPT, Gemini, and DeepSeek have revolutionized how students interact with textual knowledge by enabling natural language queries over complex material.")
    h2(doc, "1.2 Voice-Controlled Systems Overview")
    para(doc, "Voice-controlled systems allow users to interact with software through spoken commands rather than traditional keyboard input. These systems combine Speech-to-Text (STT) conversion, natural language understanding, and Text-to-Speech (TTS) synthesis. Modern browsers implement the Web Speech API, enabling real-time voice input without native desktop applications. Voice interfaces improve accessibility and support hands-free learning during revision sessions.")
    h2(doc, "1.3 Need for AI-Based Study Assistants")
    para(doc, "Students frequently manage multiple subjects, PDFs, lecture notes, and revision schedules simultaneously. Existing tools are fragmented: note apps store content, PDF readers display documents, and chatbots answer questions without persistent source context. An integrated AI study assistant addresses this gap by combining source management, grounded chat, voice interaction, and automated study artifact generation in one workspace.")
    h2(doc, "1.4 Problem Statement")
    para(doc, "The problem addressed in this project is the lack of an affordable, deployable, and feature-rich AI study platform that: (1) grounds AI responses in user-uploaded sources; (2) supports voice input and multilingual output; (3) generates diverse study materials (mind maps, flashcards, podcasts); (4) persists user notebooks across sessions; and (5) isolates data between multiple users sharing the same deployment link.")
    h2(doc, "1.5 Motivation of the Project")
    para(doc, "Commercial platforms such as Google NotebookLM demonstrate the value of source-grounded AI for learning, but are not always accessible for academic experimentation or local customization. This project was motivated by the need to build a college-level, deployable alternative using open API gateways (OpenRouter), modern web technologies, and a NotebookLM-inspired user experience.")
    h2(doc, "1.6 Objectives of the Project")
    for o in [
        "To design and implement a web-based AI study assistant with NotebookLM-style three-panel layout.",
        "To integrate voice input (STT) and voice output (TTS) for hands-free learning.",
        "To enable PDF/text source upload and source-grounded AI chat with streaming responses.",
        "To implement Studio features: mind maps, slides, flashcards, FAQ, briefing, outline, data tables, study resources, and infographics.",
        "To provide audio overview podcasts with interactive join-conversation and continue-in-podcast modes.",
        "To implement user authentication with per-user data isolation.",
        f"To deploy the application on Vercel at {LIVE_URL}.",
    ]:
        bullet(doc, o)
    h2(doc, "1.7 Scope of the Project")
    para(doc, "The scope includes frontend UI, backend API proxy, OpenRouter LLM integration, authentication, local persistence, and cloud deployment. Custom LLM training, vector database RAG pipelines, and native mobile applications are out of scope.")
    h2(doc, "1.8 Proposed Solution")
    para(doc, "The proposed Voice-Controlled AI Study Assistant is a full-stack web application comprising: (a) Vanilla JavaScript frontend with Sources | Chat | Studio panels; (b) Node.js/Express backend proxying OpenRouter API calls; (c) SSE-based streaming chat; (d) JWT stateless authentication with HMAC-derived user IDs; (e) localStorage-based notebook persistence scoped per user; (f) Vercel serverless deployment.")
    h2(doc, "1.9 Applications of the System")
    for a in ["Engineering and degree exam preparation", "Competitive exam revision (GATE, placements)", "Research paper summarization", "Language learning with multilingual TTS", "Classroom demonstration of AI in education", "Remote mentoring with shared Vercel link"]:
        bullet(doc, a)
    h2(doc, "1.10 Advantages of the Proposed System")
    for a in ["Unified workspace for sources, chat, and generated study content", "Voice and text multimodal interaction", "Low deployment cost using Vercel and OpenRouter pay-per-use", "No installation required — browser-based access", "Rich studio outputs beyond plain chat", "Per-user data isolation for shared deployments"]:
        bullet(doc, a)
    h2(doc, "1.11 Organization of the Report")
    para(doc, "Chapter 2 reviews voice assistants, AI in education, and comparative systems (Table 2.1). Chapter 3 analyses functional and non-functional requirements (Tables 3.1–3.4) and presents the system architecture (Fig. 3.1). Chapter 4 documents the complete design methodology with architecture, DFD, UML diagrams, algorithms, and security. Chapter 5 describes implementation and technology stack (Table 5.1). Chapter 6 reports experimental results (Tables 6.1–6.2). Chapter 7 concludes with achievements and future scope.")
    pb(doc)


def build_chapter2(doc):
    h1(doc, "CHAPTER 2")
    h1(doc, "LITERATURE SURVEY")
    h2(doc, "2.1 Introduction")
    para(doc, "This chapter reviews existing voice assistants, AI in education, speech technologies, NLP approaches, and comparative systems relevant to the Voice-Controlled AI Study Assistant.")
    h2(doc, "2.2 Existing Voice Assistant Technologies")
    para(doc, "Commercial voice assistants such as Google Assistant, Amazon Alexa, Apple Siri, and Microsoft Cortana demonstrate robust STT and TTS pipelines integrated with cloud NLP services. For web applications, the Web Speech API provides browser-native speech recognition and synthesis. Backend TTS services such as Google Gemini TTS offer higher quality multilingual output than browser defaults.")
    h2(doc, "2.3 AI in Smart Education Systems")
    para(doc, "Intelligent tutoring systems, adaptive learning platforms, and AI note-taking tools show growing adoption. Google NotebookLM (2023–2024) introduced source-grounded notebooks with audio overviews, marking a shift from generic chatbots to document-centric learning assistants.")
    h2(doc, "2.4 Speech Recognition Techniques")
    para(doc, "Speech recognition evolved from Hidden Markov Models (HMM) to deep learning architectures including Transformer-based models such as Whisper. Browser STT uses cloud-backed recognition engines via Web Speech API. Accuracy depends on microphone quality, accent, background noise, and language support.")
    h2(doc, "2.5 Natural Language Processing Approaches")
    para(doc, "Modern NLP for study assistants relies on LLMs for generation, summarization, translation, and structured output (JSON for mind maps, flashcards). Prompt engineering with source context replaces traditional pipeline NLP for many educational tasks. Streaming token generation via SSE improves perceived response latency.")
    h2(doc, "2.6 Machine Learning Models Used in Assistants")
    para(doc, "This project uses DeepSeek v3.2 as the primary LLM and Google Gemini 2.0 Flash as fallback, accessed through OpenRouter. TTS uses google/gemini-3.1-flash-tts-preview. No custom model training is performed; the system leverages pre-trained models via API.")
    h2(doc, "2.7 Comparative Analysis of Existing Systems")
    para(doc, "Table 2.1 compares representative study-assistant platforms across source grounding, voice support, study artifact generation, and custom deployment capability. Only the proposed system combines full voice interaction, rich studio outputs, and self-hosted Vercel deployment suitable for academic demonstration.")
    table_caption(doc, "Table 2.1 — Comparative Analysis of Existing Systems")
    tbl(doc, ["System", "Source Grounding", "Voice", "Study Artifacts", "Custom Deploy"], [
        ["Google NotebookLM", "Yes", "Limited", "Audio, slides, etc.", "No"],
        ["ChatGPT", "Via uploads", "Yes (Plus)", "Generic", "No"],
        ["Notion AI", "Workspace", "No", "Notes only", "No"],
        ["Proposed System", "Yes", "Yes (STT+TTS)", "Mind map, flashcards, podcast", "Yes (Vercel)"],
    ])
    h2(doc, "2.8 Research Gaps Identified")
    for g in ["Lack of open, deployable NotebookLM-style platforms for academic projects", "Limited integration of voice + studio artifacts in single UI", "Need for lightweight auth without expensive database for small user bases", "Multilingual learning support in Indian languages (Telugu, Hindi, Tamil)"]:
        bullet(doc, g)
    h2(doc, "2.9 Summary")
    para(doc, "Literature review confirms demand for integrated AI study platforms. The proposed system fills the gaps identified in Section 2.8 by combining NotebookLM-inspired UX, voice control, rich studio outputs, and practical Vercel deployment — substantiated by the comparison in Table 2.1 and validated in Chapter 6.")
    pb(doc)


def build_chapter3(doc, diagrams):
    h1(doc, "CHAPTER 3")
    h1(doc, "SYSTEM ANALYSIS")
    h2(doc, "3.1 Existing System")
    para(doc, "Students typically use separate tools: PDF readers, handwritten/typed notes, YouTube lectures, generic ChatGPT queries, and manual flashcard apps. These systems do not maintain persistent source-notebook relationships or generate structured study artifacts from the same uploaded material. The proposed integrated architecture that overcomes these limitations is overviewed in Fig. 3.1.")
    fig(doc, "Fig. 3.1 — System Architecture Overview", os.path.basename(diagrams["fig_3_1"]), 6.0)
    h2(doc, "3.2 Drawbacks of Existing System")
    for d in ["Fragmented workflow across multiple apps", "AI responses not consistently tied to uploaded syllabus PDFs", "No persistent notebook per subject", "Limited voice-based revision", "No integrated podcast-style audio overview from notes", "Generic chatbots lack mind maps, slide decks, and infographics"]:
        bullet(doc, d)
    h2(doc, "3.3 Proposed System")
    para(doc, f"The proposed system is a browser-based Voice-Controlled AI Study Assistant with authentication, multi-notebook management, source upload (PDF/text), streaming AI chat, Studio generation panel, audio overview, explain/translate/speak actions, and cloud deployment at {LIVE_URL}.")
    h2(doc, "3.4 Feasibility Study")
    h3(doc, "3.4.1 Technical Feasibility")
    para(doc, "Technically feasible using HTML/CSS/JavaScript frontend, Node.js/Express backend, OpenRouter API for LLM/TTS, Web Speech API for STT, pdf.js for PDF extraction, and Vercel for serverless hosting.")
    h3(doc, "3.4.2 Economic Feasibility")
    para(doc, "Vercel Hobby tier supports free deployment for academic use. OpenRouter charges per token — DeepSeek v3.2 is cost-effective for student-scale usage. No paid database is required due to localStorage and stateless auth design.")
    h3(doc, "3.4.3 Operational Feasibility")
    para(doc, "Users need only a modern browser and internet connection. Sign-up/sign-in flow is simple. Mentor login supports instructor access. Operational complexity is low for end users.")
    h2(doc, "3.5 System Requirements")
    h3(doc, "3.5.1 Hardware Requirements")
    para(doc, "Minimum client-side hardware specifications required to run the Voice-Controlled AI Study Assistant are listed in Table 3.1.")
    table_caption(doc, "Table 3.1 — Hardware Requirements")
    tbl(doc, ["Component", "Minimum Specification"], [
        ["Client PC/Laptop", "Intel i3 / equivalent, 4 GB RAM"],
        ["Microphone", "Built-in or external (for voice input)"],
        ["Speaker/Headphones", "For TTS and audio overview"],
        ["Network", "Broadband internet (≥ 2 Mbps)"],
        ["Server", "Vercel serverless (cloud — no local server needed)"],
    ])
    h3(doc, "3.5.2 Software Requirements")
    para(doc, "Table 3.2 enumerates the software components, versions, and purposes needed for development, testing, and production deployment.")
    table_caption(doc, "Table 3.2 — Software Requirements")
    tbl(doc, ["Software", "Version/Purpose"], [
        ["Operating System", "Windows 10/11, macOS, or Linux"],
        ["Web Browser", "Chrome 90+, Edge 90+, Firefox 88+"],
        ["Node.js (development)", "18.x or higher"],
        ["OpenRouter API Key", "LLM and TTS access"],
        ["Vercel CLI", "Cloud deployment"],
    ])
    h2(doc, "3.6 Functional Requirements")
    para(doc, "Ten functional requirements trace from student learning needs to implemented modules. Each requirement is verified in the test cases documented in Table 6.1.")
    table_caption(doc, "Table 3.3 — Functional Requirements")
    tbl(doc, ["ID", "Requirement"], [
        ["FR-01", "User shall register and sign in with email/password"],
        ["FR-02", "User shall create, rename, and delete notebooks"],
        ["FR-03", "User shall upload PDF/text sources"],
        ["FR-04", "User shall chat with AI grounded on enabled sources"],
        ["FR-05", "User shall use voice input for queries"],
        ["FR-06", "User shall generate studio outputs (mind map, slides, etc.)"],
        ["FR-07", "User shall listen to audio overview and join conversation"],
        ["FR-08", "User shall speak, translate, and explain AI responses"],
        ["FR-09", "User shall access saved studio outputs from Saved panel"],
        ["FR-10", "Mentor shall sign in with dedicated credentials"],
    ])
    h2(doc, "3.7 Non-Functional Requirements")
    para(doc, "Quality attributes and measurable performance targets are summarised in Table 3.4. Outcomes against these targets are evaluated in Chapter 6.")
    table_caption(doc, "Table 3.4 — Non-Functional Requirements")
    tbl(doc, ["ID", "Requirement"], [
        ["NFR-01", "Chat streaming shall begin within 5 seconds on average"],
        ["NFR-02", "UI shall be responsive on 1366×768 and above"],
        ["NFR-03", "User data shall be isolated per account"],
        ["NFR-04", "System shall support dark/light themes"],
        ["NFR-05", "Application shall be deployable on Vercel"],
        ["NFR-06", "API keys shall not be exposed in frontend code"],
    ])
    h2(doc, "3.8 System Architecture Overview")
    para(doc, "The system follows a three-tier architecture: (1) Presentation Tier — HTML/CSS/JS frontend with Web Speech API; (2) Application Tier — Express.js API on Vercel serverless functions; (3) External Services Tier — OpenRouter (DeepSeek, Gemini) for LLM and TTS. Data persistence uses browser localStorage keyed by user ID, as illustrated in Fig. 3.1.")
    pb(doc)


def build_chapter4(doc, diagrams):
    helpers = {
        "h1": h1, "h2": h2, "h3": h3, "para": para, "bullet": bullet,
        "fig": fig, "table_caption": table_caption, "tbl": tbl, "pb": pb,
    }
    build_chapter4_detailed(doc, diagrams, helpers)


def build_chapter5(doc):
    h1(doc, "CHAPTER 5")
    h1(doc, "IMPLEMENTATION")
    h2(doc, "5.1 Introduction")
    para(doc, "This chapter describes the development environment, technology stack, module implementation, and integration details of the Voice-Controlled AI Study Assistant. The authentication interface is shown in Fig. 5.1.")
    fig(doc, "Fig. 5.1 — Authentication Page", "fig_auth.png")
    h2(doc, "5.2 Development Environment")
    para(doc, f"Development was performed on Windows with VS Code/Cursor IDE. Local backend runs on Node.js port 3000. Production deployment via Vercel CLI. Live URL: {LIVE_URL}. The main three-panel interface after login is presented in Fig. 5.2.")
    fig(doc, "Fig. 5.2 — Main Three-Panel Interface (Sources | Chat | Studio)", "fig_main_ui.png")
    h2(doc, "5.3 Tools and Technologies Used")
    para(doc, "The complete technology stack — frontend libraries, backend runtime, AI models, and deployment platform — is documented in Table 5.1.")
    table_caption(doc, "Table 5.1 — Technology Stack")
    tbl(doc, ["Category", "Technology", "Purpose"], [
        ["Frontend", "HTML5, CSS3, JavaScript ES2022", "UI and client logic"],
        ["Markdown", "marked.js + highlight.js", "Render AI responses"],
        ["PDF", "pdf.js (Mozilla)", "Extract text from PDFs"],
        ["Mind Map", "Custom SVG + Reingold-Tilford", "Interactive concept maps"],
        ["Backend", "Node.js 18+, Express.js", "REST/SSE API"],
        ["AI Gateway", "OpenRouter API", "LLM and TTS access"],
        ["LLM", "deepseek/deepseek-v3.2", "Primary chat model"],
        ["TTS", "google/gemini-3.1-flash-tts-preview", "Multilingual speech"],
        ["Deployment", "Vercel", "Static + serverless hosting"],
    ])
    para(doc, "The chat interface with saved studio outputs panel is illustrated in Fig. 5.3.")
    fig(doc, "Fig. 5.3 — Chat Interface with Saved Studio Outputs", "fig_chat_saved.png")
    h2(doc, "5.4 Coding Methodology")
    para(doc, "Agile iterative development was followed. Features were added incrementally: core chat, studio outputs, voice, authentication, and deployment fixes. Single-page application pattern with modular functions in app.js (~2800 lines) and backend consolidated in server.js (~880 lines). An interactive mind map output is shown in Fig. 5.4.")
    fig(doc, "Fig. 5.4 — Interactive Mind Map Studio Output", "fig_mindmap.png")
    h2(doc, "5.5 Implementation of Modules")
    para(doc, "Key files: frontend/index.html, frontend/app.js, frontend/styles.css, frontend/auth.html, frontend/mindmap.js, backend/server.js, vercel.json.")
    h2(doc, "5.6 Integration of Voice Commands")
    para(doc, "Mic button triggers SpeechRecognition. Final transcript is passed to sendMessage(). TTS via POST /api/tts returns WAV audio from Gemini TTS. speakText(), playExplainAudio(), and podcast modes use pipelined audio playback.")
    h2(doc, "5.7 AI Model Training and Testing")
    para(doc, "No custom training was performed. Pre-trained models are accessed via OpenRouter with DeepSeek v3.2 primary and Gemini 2.0 Flash fallback. Prompt engineering ensures source grounding and structured JSON for studio endpoints.")
    h2(doc, "5.8 User Authentication Process")
    para(doc, "Stateless auth: uid = HMAC(email + password). JWT stored in localStorage under sa_auth. Mentor login via username ss6156. Auth guard redirects unauthenticated users to auth.html.")
    h2(doc, "5.9 API Integration")
    para(doc, "Major endpoints: /api/auth/signup, /api/auth/signin, /api/chat (SSE), /api/tts, /api/explain, /api/podcast, /api/studio/mindmap, /api/studio/slides, /api/studio/flashcards, /api/studio/faq, /api/studio/briefing, /api/studio/outline, /api/studio/datatable, /api/studio/resources, /api/studio/infographic.")
    h2(doc, "5.10 Screenshots of Implementation")
    para(doc, "Implementation screenshots are presented in their respective sections: Fig. 5.1 (Section 5.1), Fig. 5.2 (Section 5.2), Fig. 5.3 (Section 5.3), and Fig. 5.4 (Section 5.4).")
    pb(doc)


def build_chapter6(doc):
    h1(doc, "CHAPTER 6")
    h1(doc, "RESULTS AND DISCUSSION")
    h2(doc, "6.1 Experimental Setup")
    para(doc, f"Testing was performed on local development (localhost:3000), production deployment ({LIVE_URL}), Chrome and Edge browsers, multiple user accounts, and PDF sources up to 50 pages. A sample chat response with action buttons is shown in Fig. 6.1.")
    fig(doc, "Fig. 6.1 — Chat Response with Action Buttons", "fig_chat_actions.png")
    h2(doc, "6.2 Test Cases")
    para(doc, "Twelve test cases were designed to verify every functional requirement in Table 3.3. Execution was performed on both localhost and the Vercel production deployment. Results are recorded in Table 6.1.")
    table_caption(doc, "Table 6.1 — Test Cases and Results")
    tbl(doc, ["TC ID", "Test Case", "Expected Result", "Status"], [
        ["TC-01", "User sign-up with valid email", "JWT returned, redirect to app", "Pass"],
        ["TC-02", "User sign-in with correct password", "Same uid, notebooks loaded", "Pass"],
        ["TC-03", "Upload PDF source", "Text extracted, listed in Sources", "Pass"],
        ["TC-04", "Chat query on enabled source", "Streaming AI response", "Pass"],
        ["TC-05", "Voice input query", "Transcript sent, response received", "Pass"],
        ["TC-06", "Generate mind map", "Interactive SVG displayed and saved", "Pass"],
        ["TC-07", "Audio overview generation", "Host-expert dialogue with TTS", "Pass"],
        ["TC-08", "Explain button after response", "Audio explanation plays", "Pass"],
        ["TC-09", "Translate to Telugu", "Translated text + TTS", "Pass"],
        ["TC-10", "Cross-device Vercel access", "App loads; API responds", "Pass"],
        ["TC-11", "Mentor login (ss6156)", "Mentor role badge displayed", "Pass"],
        ["TC-12", "Action buttons after stream", "Speak/Explain visible", "Pass"],
    ])
    para(doc, "A study resources roadmap generated from uploaded sources is presented in Fig. 6.2.")
    fig(doc, "Fig. 6.2 — Study Resources Output", "fig_study_resources.png")
    h2(doc, "6.3 Performance Evaluation")
    para(doc, "Overall system performance was assessed against the non-functional targets in Table 3.4. Average chat first-token latency is 2–5 seconds. Full mind map generation takes 15–30 seconds. TTS per chunk takes 1–3 seconds. Vercel cold start may add ~6 seconds for initial auth requests.")
    h2(doc, "6.4 Accuracy of Speech Recognition")
    para(doc, "Web Speech API accuracy is high for clear English speech (~90%+ in quiet environments). Indian accent and Telugu/Hindi mixed speech show moderate accuracy depending on browser engine. Text input remains a reliable fallback.")
    h2(doc, "6.5 Response Time Analysis")
    para(doc, "Detailed response time measurements for authentication, chat streaming, studio generation, TTS, and PDF extraction are summarised in Table 6.2.")
    table_caption(doc, "Table 6.2 — Performance Metrics")
    tbl(doc, ["Operation", "Average Time"], [
        ["Sign-in API", "~6 s (cold) / ~1 s (warm)"],
        ["Chat first token (SSE)", "2–5 s"],
        ["Full chat response (200 words)", "8–15 s"],
        ["Mind map generation", "15–30 s"],
        ["TTS single chunk", "1–3 s"],
        ["PDF text extraction (20 pages)", "3–8 s (client-side)"],
    ])
    h2(doc, "6.6 Comparative Performance Analysis")
    para(doc, "Compared to a manual study workflow, the system reduces time to create flashcards, mind maps, and summaries from hours to minutes. Compared to generic ChatGPT, responses are better grounded when sources are uploaded. Table 2.1 provides the qualitative platform comparison; Table 6.2 quantifies the operational response times measured in this project.")
    h2(doc, "6.7 Output Screenshots")
    para(doc, "Output screenshots are included in Fig. 6.1 (Section 6.1) and Fig. 6.2 (Section 6.2).")
    h2(doc, "6.8 Discussion of Results")
    para(doc, f"The system successfully demonstrates a NotebookLM-style learning workflow with voice control at {LIVE_URL}. Uploading programming notes and querying concepts such as pointer-to-pointer produced accurate C/C++ explanations with saved mind maps, flashcards, slides, and audio outputs.")
    h2(doc, "6.9 Limitations of the System")
    for lim in ["Data stored in browser localStorage — not synced across devices", "No vector RAG — large PDFs may exceed context limits", "Stateless auth — no password reset or email verification", "Serverless cold starts may cause initial delay", "STT quality varies by browser and accent"]:
        bullet(doc, lim)
    pb(doc)


def build_chapter7_and_back(doc):
    h1(doc, "CHAPTER 7")
    h1(doc, "CONCLUSION AND FUTURE SCOPE")
    h2(doc, "7.1 Conclusion")
    para(doc, f"The Voice-Controlled AI Study Assistant successfully delivers an integrated, deployable, AI-powered learning platform at {LIVE_URL}, combining source-grounded chat, voice interaction, studio artifact generation, audio podcasts, and user authentication.")
    h2(doc, "7.2 Achievements of the Project")
    for a in ["Designed and implemented full-stack NotebookLM-style UI", "Integrated DeepSeek v3.2 via OpenRouter with SSE streaming", "Built 10+ studio features including mind map, slides, flashcards, infographic", "Implemented voice input, TTS output, audio overview, and explain modes", "Added JWT authentication with per-user data isolation", "Deployed production application on Vercel"]:
        bullet(doc, a)
    h2(doc, "7.3 Future Enhancements")
    for f in ["Cloud database (Supabase) for cross-device notebook sync", "Vector RAG with chunking for large documents", "Offline PWA support with service workers", "Collaborative notebooks for group study", "Analytics dashboard for mentors"]:
        bullet(doc, f)
    h2(doc, "7.4 Real-Time Industrial Applications")
    para(doc, "EdTech startups, coaching institutes, corporate L&D departments, and library digitization projects can adapt this architecture for AI tutoring and training content generation with minimal infrastructure cost.")
    h2(doc, "7.5 Final Remarks")
    para(doc, "This project demonstrates that modern LLM APIs, browser speech technologies, and serverless deployment can be combined to build sophisticated educational tools without massive infrastructure.")
    pb(doc)

    h1(doc, "REFERENCES")
    refs = [
        '[1] Google, "NotebookLM: An AI-powered research and writing assistant," Google Labs, 2023–2024.',
        "[2] OpenRouter, OpenRouter API Documentation, https://openrouter.ai/docs, 2024.",
        '[3] DeepSeek, "DeepSeek-V3 Technical Report," 2024.',
        "[4] Mozilla, pdf.js Documentation, https://mozilla.github.io/pdf.js/, 2024.",
        "[5] W3C, Web Speech API Specification, https://w3c.github.io/speech-api/, 2024.",
        "[6] AntV, Infographic Library, https://github.com/antvis/Infographic, 2024.",
        "[7] Vercel, Serverless Functions Documentation, https://vercel.com/docs/functions, 2024.",
        "[8] IETF RFC 7519, JSON Web Token (JWT), 2015.",
        "[9] Russell, S., and Norvig, P., Artificial Intelligence: A Modern Approach, 4th ed., Pearson, 2020.",
        "[10] Jurafsky, D., and Martin, J. H., Speech and Language Processing, 3rd ed. draft, 2023.",
    ]
    for r in refs:
        para(doc, r, align=WD_ALIGN_PARAGRAPH.LEFT)
    pb(doc)

    h1(doc, "APPENDICES")
    h2(doc, "Appendix A – Source Code")
    para(doc, "voice-study-assistant/\n├── frontend/ (index.html, auth.html, app.js, styles.css, mindmap.js)\n├── backend/ (server.js, package.json)\n├── vercel.json\n└── README.md")
    h2(doc, "Appendix B – Sample Voice Commands")
    tbl(doc, ["Voice Command", "Action"], [
        ["Explain photosynthesis", "Sends query to AI chat"],
        ["What is a dangling pointer?", "C programming query from PDF source"],
        ["Summarize chapter 3", "Summary from uploaded source"],
        ["(Mic + Telugu speech)", "Telugu query with Telugu response"],
    ])
    h2(doc, "Appendix C – Test Case Reports")
    para(doc, "Detailed test execution logs with date, tester name, browser, and pass/fail status are maintained in the project logbook. A summary of all twelve test cases appears in Table 6.1.")
    h2(doc, "Appendix D – User Manual")
    para(doc, f"1. Open {LIVE_URL}\n2. Sign up with email and password\n3. Create a notebook and upload PDF/text sources\n4. Ask questions via text or microphone\n5. Use Studio cards to generate mind maps, slides, flashcards\n6. Click Speak/Explain/Translate on AI responses\n7. Access saved outputs in Saved panel\n8. Sign out via avatar")
    h2(doc, "Appendix E – Installation Guide")
    para(doc, "Local: cd backend && npm install && node server.js. Deploy: vercel --prod with OPENROUTER_API_KEY set.")
    h2(doc, "Appendix F – Dataset Details and Screenshots")
    para(doc, "No custom ML dataset was created. Testing used publicly available study PDFs and user-entered topics. Implementation and output screenshots referenced in this report are Fig. 5.1–5.4 (Chapter 5) and Fig. 6.1–6.2 (Chapter 6); these figures are embedded in the main body and do not require a separate appendix.")

    h1(doc, "ANNEXURES")
    for a in [
        "Annexure I – Project Synopsis",
        "Annexure II – Weekly Progress Reports",
        "Annexure III – Gantt Chart / Project Schedule",
        "Annexure IV – Plagiarism Report",
        "Annexure V – Guide Meeting Logbook",
        "Annexure VI – Student Contribution Sheet",
        "Annexure VII – Attendance Record",
        "Annexure VIII – Ethical Clearance (if applicable)",
        "Annexure IX – Publication/Patent Details (if any)",
        "Annexure X – Project Completion Certificate",
    ]:
        para(doc, a, bold=True)
        para(doc, "[Attach signed/scanned document as per department format]")


def build():
    os.makedirs(ASSETS, exist_ok=True)
    print("Generating Graphviz diagrams...")
    diagrams = generate_all()
    if capture_screenshots and not os.path.isfile(os.path.join(ASSETS, "fig_auth.png")):
        print("Capturing screenshots...")
        capture_screenshots()

    doc = Document()
    # Default section — cover + bonafide (roman page i)
    _set_margins(doc.sections[0])
    _add_page_number_footer(doc.sections[0], "lowerRoman", start=1)

    build_cover_and_bonafide(doc)

    # Preliminary pages — roman numerals continue (approval = ii, …)
    new_section(doc, "lowerRoman")

    build_preliminary(doc)

    # Chapters — arabic numerals from 1
    new_section(doc, "decimal", start=1)

    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc, diagrams)
    build_chapter4(doc, diagrams)
    build_chapter5(doc)
    build_chapter6(doc)
    build_chapter7_and_back(doc)

    saved_path = None
    for out_path in (OUTPUT, OUTPUT_ALT):
        try:
            doc.save(out_path)
            saved_path = out_path
            print(f"Report saved: {out_path}")
            break
        except PermissionError:
            print(f"Could not write {out_path} (file may be open in Word)")
    else:
        raise PermissionError("Could not save report — close the DOCX in Word and re-run build_major_report.py")

    try:
        from update_toc_pages import update_toc
        update_toc(saved_path)
        print("TOC page numbers synced from Word layout")
    except Exception as exc:
        print(f"TOC auto-sync skipped: {exc}")


if __name__ == "__main__":
    build()
